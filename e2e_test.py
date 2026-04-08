#!/usr/bin/env python3
"""
Full End-to-End Test with Content Validation
Tests the entire CV generation pipeline and validates output quality
"""

import os
import sys
import json
from pathlib import Path
from datetime import datetime

sys.path.insert(0, str(Path(__file__).parent))

from dotenv import load_dotenv
load_dotenv()

# Test job description (realistic)
TEST_JOB = """
Senior Software Engineer - Python/AWS
Company: Acme Technologies Ltd
Location: London, UK (Hybrid - 2 days in office)

About Us:
Acme Technologies is a fast-growing fintech startup revolutionizing payment processing.

The Role:
We're looking for a talented Software Engineer to join our platform team. You'll be building 
scalable microservices that handle millions of transactions daily.

Requirements:
- 3+ years of Python development experience
- Strong experience with AWS (Lambda, S3, DynamoDB, SQS)
- Experience with Docker and Kubernetes
- PostgreSQL or similar relational databases
- RESTful API design and implementation
- CI/CD pipelines (GitHub Actions, Jenkins)
- Agile/Scrum methodology

Nice to Have:
- Experience with Apache Kafka or similar messaging systems
- Terraform or CloudFormation for infrastructure as code
- Machine Learning/AI experience
- Financial services background

What We Offer:
- Competitive salary £70-90k
- Stock options
- Flexible working
- Learning budget
"""

def validate_cv_content(cv_data: dict, job_info: dict) -> dict:
    """Validate CV content quality"""
    issues = []
    warnings = []
    
    # Check bio
    bio = cv_data.get('bio', '')
    if not bio:
        issues.append("Missing bio")
    elif len(bio) < 100:
        warnings.append(f"Bio too short ({len(bio)} chars)")
    elif len(bio) > 400:
        warnings.append(f"Bio too long ({len(bio)} chars) - may overflow page")
    
    # Check for forbidden words in bio
    forbidden = ['senior', 'lead', 'principal', 'staff']
    bio_lower = bio.lower()
    for word in forbidden:
        if word in bio_lower:
            issues.append(f"Bio contains forbidden word: '{word}'")
    
    # Check expertise
    expertise = cv_data.get('expertise', [])
    if not expertise:
        issues.append("Missing expertise list")
    elif len(expertise) < 10:
        warnings.append(f"Only {len(expertise)} skills (expected 12-14)")
    elif len(expertise) > 16:
        warnings.append(f"Too many skills ({len(expertise)}) - may overflow")
    
    # Check experience sections
    for section_key, section_name in [('t', 'T. Rowe Price'), ('a', 'AWS')]:
        section = cv_data.get(section_key, {})
        
        if not section:
            issues.append(f"Missing {section_name} section")
            continue
        
        skills = section.get('skills', '')
        if not skills:
            warnings.append(f"{section_name}: Missing tech stack")
        
        # Check bullet points
        bp_count = 0
        for i in range(1, 5):
            bp = section.get(f'bp{i}', '')
            if bp:
                bp_count += 1
                if len(bp) < 50:
                    warnings.append(f"{section_name} bp{i}: Too short ({len(bp)} chars)")
                elif len(bp) > 400:
                    warnings.append(f"{section_name} bp{i}: Too long ({len(bp)} chars)")
                
                # Check for metrics/numbers
                import re
                if not re.search(r'\d', bp):
                    warnings.append(f"{section_name} bp{i}: No quantified metrics")
        
        if bp_count < 3:
            issues.append(f"{section_name}: Only {bp_count} bullet points (need 3-4)")
    
    # Check for job-relevant keywords
    job_skills = job_info.get('required_skills', []) + job_info.get('preferred_skills', [])
    cv_text = json.dumps(cv_data).lower()
    matched_skills = [s for s in job_skills if s.lower() in cv_text]
    match_rate = len(matched_skills) / len(job_skills) * 100 if job_skills else 0
    
    if match_rate < 30:
        issues.append(f"Low skill match rate: {match_rate:.0f}%")
    elif match_rate < 50:
        warnings.append(f"Moderate skill match rate: {match_rate:.0f}%")
    
    return {
        'valid': len(issues) == 0,
        'issues': issues,
        'warnings': warnings,
        'skill_match_rate': match_rate,
        'bio_length': len(bio),
        'expertise_count': len(expertise),
    }

def validate_docx_content(docx_path: str) -> dict:
    """Validate the actual DOCX file content"""
    from docx import Document
    
    if not Path(docx_path).exists():
        return {'valid': False, 'error': 'File does not exist'}
    
    try:
        doc = Document(docx_path)
        
        # Extract all text
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        full_text.append(para.text)
        
        text = '\n'.join(full_text)
        
        # Check for unfilled placeholders
        import re
        placeholders = re.findall(r'\{\{[^}]+\}\}', text)
        
        # Check metadata
        props = doc.core_properties
        
        return {
            'valid': len(placeholders) == 0,
            'unfilled_placeholders': placeholders,
            'word_count': len(text.split()),
            'char_count': len(text),
            'author': props.author,
            'title': props.title,
            'has_ai_indicators': 'ai' in (props.comments or '').lower() or 'claude' in (props.comments or '').lower(),
        }
        
    except Exception as e:
        return {'valid': False, 'error': str(e)}

def run_full_e2e_test():
    """Run complete end-to-end test with content validation"""
    print("=" * 70)
    print("FULL END-TO-END TEST WITH CONTENT VALIDATION")
    print("=" * 70)
    
    api_key = os.getenv('ANTHROPIC_API_KEY')
    if not api_key:
        print("✗ No API key configured")
        return False
    
    from web_app import BatchCVGenerator, CV_VARIANTS
    from ats_scorer import ATSScorer
    
    generator = BatchCVGenerator(api_key)
    scorer = ATSScorer()
    
    # Step 1: Parse job description
    print("\n1. PARSING JOB DESCRIPTION")
    print("-" * 40)
    
    job_info = generator.parse_job_description(TEST_JOB)
    # Store raw text for ATS phrase extraction
    job_info['raw_text'] = TEST_JOB
    print(f"   Job Title: {job_info.get('job_title')}")
    print(f"   Company: {job_info.get('company_name')}")
    print(f"   Location: {job_info.get('location')}")
    print(f"   Required Skills: {', '.join(job_info.get('required_skills', []))}")
    print(f"   Preferred Skills: {', '.join(job_info.get('preferred_skills', []))}")
    
    if not job_info.get('job_title') or job_info.get('job_title') == 'Unknown':
        print("   ✗ Failed to parse job title")
        return False
    print("   ✓ Job parsing successful")
    
    # Step 2: Generate all variants
    print("\n2. GENERATING CV VARIANTS")
    print("-" * 40)
    
    test_output = Path('outputs/e2e_test')
    test_output.mkdir(parents=True, exist_ok=True)
    
    all_results = []
    
    for variant_key in CV_VARIANTS.keys():
        print(f"\n   [{variant_key.upper()}]")
        
        try:
            generated = generator.generate_cv_and_cover_letter(job_info, variant_key)
            cv_data = generated.get('cv', {})
            
            # Apply LLM enhancement (same as web app does)
            from web_app import llm_enhance_cv_content
            cv_data = llm_enhance_cv_content(cv_data, job_info, variant_key)
            
            # Validate CV content
            validation = validate_cv_content(cv_data, job_info)
            
            print(f"   Bio: {cv_data.get('bio', '')[:60]}...")
            print(f"   Expertise: {len(cv_data.get('expertise', []))} skills")
            print(f"   Skill Match: {validation['skill_match_rate']:.0f}%")
            
            if validation['issues']:
                for issue in validation['issues']:
                    print(f"   ✗ ISSUE: {issue}")
            
            if validation['warnings']:
                for warn in validation['warnings'][:3]:
                    print(f"   ⚠ Warning: {warn}")
            
            # Create DOCX
            docx_path = test_output / f"CV_{variant_key}.docx"
            pdf_path = generator.create_cv_docx(cv_data, job_info, str(docx_path), create_pdf=True)
            
            # Validate DOCX
            docx_validation = validate_docx_content(str(docx_path))
            
            if docx_path.exists():
                print(f"   ✓ DOCX created ({docx_path.stat().st_size} bytes)")
                
                if docx_validation.get('unfilled_placeholders'):
                    print(f"   ✗ Unfilled placeholders: {docx_validation['unfilled_placeholders']}")
                else:
                    print(f"   ✓ No unfilled placeholders")
                
                if docx_validation.get('has_ai_indicators'):
                    print(f"   ✗ AI indicators in metadata")
                else:
                    print(f"   ✓ Clean metadata (author: {docx_validation.get('author')})")
            else:
                print(f"   ✗ DOCX NOT created")
            
            if pdf_path and Path(pdf_path).exists():
                print(f"   ✓ PDF created ({Path(pdf_path).stat().st_size} bytes)")
            
            # ATS Score
            cv_text = f"{cv_data.get('bio', '')} {' '.join(cv_data.get('expertise', []))}"
            job_skills = job_info.get('required_skills', []) + job_info.get('preferred_skills', [])
            ats_result = scorer.score_cv(cv_text, TEST_JOB, job_skills)
            print(f"   ATS Score: {ats_result['total_score']:.1f}/100 ({ats_result['grade']})")
            
            all_results.append({
                'variant': variant_key,
                'cv_valid': validation['valid'],
                'docx_valid': docx_validation.get('valid', False),
                'ats_score': ats_result['total_score'],
                'issues': validation['issues'],
            })
            
        except Exception as e:
            print(f"   ✗ Error: {e}")
            import traceback
            traceback.print_exc()
            all_results.append({
                'variant': variant_key,
                'error': str(e),
            })
    
    # Step 3: Generate cover letter
    print("\n3. GENERATING COVER LETTER")
    print("-" * 40)
    
    try:
        cover_letter_path = test_output / "Cover_Letter.pdf"
        best_generated = generator.generate_cv_and_cover_letter(job_info, 'professional')
        cover_letter_text = best_generated.get('cover_letter', '')
        
        print(f"   Length: {len(cover_letter_text)} chars")
        print(f"   Preview: {cover_letter_text[:100]}...")
        
        generator.create_cover_letter_pdf(cover_letter_text, job_info, str(cover_letter_path))
        
        if cover_letter_path.exists():
            print(f"   ✓ Cover letter PDF created ({cover_letter_path.stat().st_size} bytes)")
        else:
            print(f"   ✗ Cover letter NOT created")
            
    except Exception as e:
        print(f"   ✗ Error: {e}")
    
    # Summary
    print("\n" + "=" * 70)
    print("TEST SUMMARY")
    print("=" * 70)
    
    passed = sum(1 for r in all_results if r.get('cv_valid') and r.get('docx_valid'))
    failed = len(all_results) - passed
    
    print(f"\nVariants Generated: {len(all_results)}")
    print(f"Fully Valid: {passed}")
    print(f"With Issues: {failed}")
    
    for r in all_results:
        status = "✓" if r.get('cv_valid') and r.get('docx_valid') else "✗"
        score = r.get('ats_score', 0)
        print(f"  {status} {r['variant']}: ATS {score:.0f}/100", end="")
        if r.get('issues'):
            print(f" - Issues: {len(r['issues'])}")
        elif r.get('error'):
            print(f" - ERROR: {r['error']}")
        else:
            print()
    
    print(f"\nOutput folder: {test_output}")
    print("\n" + "=" * 70)
    
    return passed == len(all_results)

if __name__ == '__main__':
    success = run_full_e2e_test()
    sys.exit(0 if success else 1)
