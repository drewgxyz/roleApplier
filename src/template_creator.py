import json
import re
from pathlib import Path
from typing import Dict, List, Union
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import subprocess
import sys
from datetime import datetime
import os

class CVCustomizer:
    def __init__(self, template_path: str):
        """Initialize with a Word template containing placeholders."""
        self.template_path = template_path
        self.document = Document(template_path)
    
    def replace_placeholders(self, replacements: Dict[str, Union[str, List[str]]]):
        """Replace placeholders throughout the document while preserving formatting."""
        # Handle paragraphs
        for paragraph in self.document.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)
        
        # Handle tables
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)
        
        # Handle headers and footers
        for section in self.document.sections:
            # Header
            for paragraph in section.header.paragraphs:
                self._replace_in_paragraph(paragraph, replacements)
            # Footer
            for paragraph in section.footer.paragraphs:
                self._replace_in_paragraph(paragraph, replacements)
    
    def _replace_in_paragraph(self, paragraph, replacements):
        """Replace placeholders in a paragraph while preserving run formatting."""
        # Get full paragraph text
        full_text = paragraph.text
        
        # Check if any placeholders exist in this paragraph
        has_placeholder = any(f"{{{{{key}}}}}" in full_text for key in replacements)
        if not has_placeholder:
            return
        
        # Process each placeholder
        for key, value in replacements.items():
            placeholder = f"{{{{{key}}}}}"
            
            if placeholder in full_text:
                if isinstance(value, list):
                    # For lists, create bullet points
                    value_text = '\n• '.join(value)
                    value_text = '• ' + value_text if value else ''
                else:
                    value_text = str(value)
                
                # If it's a simple replacement (placeholder is complete in one run)
                for run in paragraph.runs:
                    if placeholder in run.text:
                        # Preserve formatting
                        run.text = run.text.replace(placeholder, value_text)
                        return
                
                # Complex case: placeholder spans multiple runs
                # Rebuild the paragraph
                self._complex_replace(paragraph, placeholder, value_text)
    
    def _complex_replace(self, paragraph, placeholder, replacement):
        """Handle placeholders that span multiple runs."""
        # Store run properties
        run_props = []
        combined_text = ""
        
        for run in paragraph.runs:
            run_props.append({
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'color': run.font.color.rgb if run.font.color.rgb else None
            })
            combined_text += run.text
        
        # Replace in combined text
        new_text = combined_text.replace(placeholder, replacement)
        
        # Clear runs and recreate with original formatting
        for run in paragraph.runs:
            run.text = ""
        
        # Add the new text with the first run's formatting
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            # Apply the original formatting from the first run
            if run_props:
                props = run_props[0]
                paragraph.runs[0].bold = props['bold']
                paragraph.runs[0].italic = props['italic']
                paragraph.runs[0].underline = props['underline']
                if props['font_name']:
                    paragraph.runs[0].font.name = props['font_name']
                if props['font_size']:
                    paragraph.runs[0].font.size = props['font_size']
    
    def save_docx(self, output_path: str):
        """Save the modified document as a Word file."""
        self.document.save(output_path)
    
    def convert_to_pdf(self, docx_path: str, pdf_path: str):
        """Convert Word document to PDF using LibreOffice (cross-platform)."""
        try:
            # Using LibreOffice in headless mode
            subprocess.run([
                'soffice',
                '--headless',
                '--convert-to',
                'pdf',
                '--outdir',
                str(Path(pdf_path).parent),
                docx_path
            ], check=True)
            
            # Rename to desired output name if needed
            generated_pdf = Path(docx_path).with_suffix('.pdf')
            if generated_pdf.name != Path(pdf_path).name:
                generated_pdf.rename(pdf_path)
                
        except subprocess.CalledProcessError:
            print("LibreOffice conversion failed. Trying python-docx2pdf...")
            # Fallback to python-docx2pdf (Windows only)
            try:
                from docx2pdf import convert
                convert(docx_path, pdf_path)
            except ImportError:
                print("Please install LibreOffice or docx2pdf for PDF conversion")
                raise
    
    def customize_cv(self, job_data: Dict, output_name: str):
        """Main method to customize CV with job data."""
        # Create execution folder with today's date and time
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        execution_folder = f"outputs/{timestamp}_{output_name}"
        
        # Create the execution folder
        os.makedirs(execution_folder, exist_ok=True)
        print(f"📁 Created execution folder: {execution_folder}")
        
        # Replace all placeholders
        self.replace_placeholders(job_data)
        
        # Save as Word document in execution folder
        docx_output = f"{execution_folder}/{output_name}.docx"
        self.save_docx(docx_output)
        print(f"✓ Word document saved: {docx_output}")
        
        # Convert to PDF in execution folder
        pdf_output = f"{execution_folder}/{output_name}.pdf"
        self.convert_to_pdf(docx_output, pdf_output)
        print(f"✓ PDF generated: {pdf_output}")
        
        return docx_output, pdf_output


def main():
    """Example usage of the CV customizer."""
    
    # Example job posting data
    job_data = {
        "job_title": "Senior Python Developer",
        "company_name": "TechCorp Solutions",
        "location": "San Francisco, CA",
        "skills": ["Python", "Django", "PostgreSQL", "Docker", "AWS"],
        "years_experience": "5+",
        "experience_highlights": [
            "Led development of microservices architecture serving 1M+ users",
            "Reduced API response time by 60% through optimization",
            "Mentored team of 4 junior developers"
        ],
        "education_focus": "Computer Science",
        "contact_email": "john.doe@email.com",
        "contact_phone": "+1-555-0123"
    }
    
    customizer = CVCustomizer("resources/template.docx")
    
    # Generate customized CV
    customizer.customize_cv(job_data, "CV_TechCorp_Senior_Python")
    
    # You can also load data from JSON file
    # with open('job_posting.json', 'r') as f:
    #     job_data = json.load(f)
    # customizer.customize_cv(job_data, "CV_CustomJob")


if __name__ == "__main__":
    main()


# === AUTOMATION SCRIPT ===
# Save this as 'batch_customize.py' for bulk processing

def batch_process_cvs(template_path: str, jobs_folder: str):
    """Process multiple job applications from JSON files."""
    jobs_path = Path(jobs_folder)
    
    for json_file in jobs_path.glob("*.json"):
        print(f"\nProcessing: {json_file.name}")
        
        with open(json_file, 'r') as f:
            job_data = json.load(f)
        
        # Create output name from company and position
        company = job_data.get('company_name', 'Unknown').replace(' ', '_')
        position = job_data.get('job_title', 'Position').replace(' ', '_')
        output_name = f"CV_{company}_{position}"
        
        # Customize CV
        customizer = CVCustomizer(template_path)
        customizer.customize_cv(job_data, output_name)


# === TEMPLATE CREATION GUIDE ===
"""
Creating Your Word Template:

1. Open Word and create your CV layout
2. Insert placeholders where you want dynamic content:
   - {{job_title}}
   - {{company_name}}
   - {{skills}}
   - {{experience_highlights}}
   - etc.

3. Format the placeholders with your desired styling
   (the script preserves formatting)

4. Save as 'cv_template.docx'

Example template structure:

---
JOHN DOE
{{job_title}} | {{location}}
Email: {{contact_email}} | Phone: {{contact_phone}}

OBJECTIVE
Seeking position as {{job_title}} at {{company_name}} where I can utilize 
my {{years_experience}} years of experience.

SKILLS
{{skills}}

EXPERIENCE HIGHLIGHTS
{{experience_highlights}}

EDUCATION
Bachelor of Science in {{education_focus}}
---
"""