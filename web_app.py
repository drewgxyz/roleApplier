"""
Flask Web Application for Batch CV and Cover Letter Generation
Uses Claude 3.5 Sonnet for ATS-friendly CV generation
With configurable skills, experiences, and settings stored in SQLite
"""

import os
import json
import re
import subprocess
from datetime import datetime
from pathlib import Path
from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for
from dotenv import load_dotenv
import anthropic
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import zipfile
import requests
from bs4 import BeautifulSoup
from urllib.parse import urlparse
from typing import List, Optional

# Import database functions
from database import (
    get_all_skills, get_approved_skills, get_blacklisted_skills,
    get_skills_by_tier, get_experiences, get_settings, update_setting,
    add_skill, update_skill, delete_skill, add_experience, update_experience,
    add_bullet_point, update_bullet_point, delete_bullet_point, get_db
)

# Import ATS scorer
from ats_scorer import ATSScorer, validate_single_page

# CV Variant styles for different tones
CV_VARIANTS = {
    'professional': {
        'name': 'Professional',
        'description': 'Formal, corporate tone. Best for traditional companies.',
        'tone_instruction': 'Use formal, professional language. Focus on achievements and responsibilities. Avoid casual phrases.',
    },
    'technical': {
        'name': 'Technical',
        'description': 'Tech-focused with detailed technical depth. Best for engineering roles.',
        'tone_instruction': 'Emphasize technical details, specific technologies, and engineering achievements. Use precise technical terminology.',
    },
    'impact': {
        'name': 'Impact-Driven',
        'description': 'Results and metrics focused. Best for startups and growth companies.',
        'tone_instruction': 'Lead with quantified results and business impact. Use action verbs and metrics prominently. Show ROI and efficiency gains.',
    },
}

# Load environment variables
load_dotenv()

app = Flask(__name__)

# Configuration
ANTHROPIC_API_KEY = os.getenv('ANTHROPIC_API_KEY')
TEMPLATE_PATH = 'resources/template.docx'

# Model configuration - use appropriate models for different tasks
# Sonnet 4: Quality writing tasks (CV generation, cover letters, job parsing)
# Haiku 3.5: Simple extraction/cleanup tasks (URL scraping, enhancement passes)
CLAUDE_MODEL_QUALITY = "claude-sonnet-4-6"  # For quality-critical tasks
CLAUDE_MODEL_FAST = "claude-haiku-4-5"  # For simpler tasks (10x cheaper)
CLAUDE_MODEL_JUDGE = "claude-opus-4-8"  # For LLM-as-judge feedback layer

# ===== TESTING FLAGS (set in .env) =====
def _env_bool(key: str, default: bool = True) -> bool:
    """Parse boolean from env var (true/false/1/0)"""
    val = os.getenv(key, str(default)).lower()
    return val in ('true', '1', 'yes')

ENABLE_PROFESSIONAL = _env_bool('ENABLE_PROFESSIONAL', False)
ENABLE_TECHNICAL = _env_bool('ENABLE_TECHNICAL', False)
ENABLE_IMPACT = _env_bool('ENABLE_IMPACT', True)
ENABLE_COVER_LETTER = _env_bool('ENABLE_COVER_LETTER', False)
ENABLE_ENHANCEMENT = _env_bool('ENABLE_ENHANCEMENT', True)
ENABLE_JUDGE = _env_bool('ENABLE_JUDGE', True)

# ===== CODE-SIDE BULLET OVERRIDES =====
# Canonical bullet descriptions keyed by (company_name, bullet_index).
# These take priority over whatever the local DB has, ensuring consistent
# metrics across machines without requiring DB updates.
# Tech placeholders ({tech}, {cloud}, etc.) are preserved and filled at runtime.
BULLET_OVERRIDES = {
    ("Compare the Market", 0): (
        "Led architecture of cross-team AI automation converting 50+ Product "
        "Requirement Documents (PRDs) into JIRA-ready tickets using {tech}, "
        "reducing planning time from 2 hours to 15 minutes - adopted by 6+ "
        "product managers including Head of Product"
    ),
    ("Compare the Market", 1): (
        "Designed and shipped a 7-agent {tech}-based AI pipeline with Redis "
        "and parallel Python workers, enabling end-to-end PRD processing with "
        "automated task routing and production-ready knowledge graph indexing "
        "on AWS EFS"
    ),
    ("Compare the Market", 2): (
        "Deployed AI-powered code review system processing 34,000+ merge "
        "requests at ~95% adoption across 400+ engineers, integrating {tech} "
        "for automated security checks, codebase-aware review suggestions, "
        "and CI/CD workflow automation"
    ),
    ("Compare the Market", 3): (
        "Mentored 100+ engineers on AI-native development and large language "
        "models through 5 workshops and 2 hackathons, supporting teams "
        "shipping production features with AI-powered Python backends"
    ),
    ("T. Rowe Price", 0): (
        "Led architecture of production-grade {tech} data migration tool "
        "syncing complex relational data across DEV/STAGE/PROD environments "
        "with rollback safety, referential integrity validation via JSON APIs, "
        "and automated pytest suites - reducing migration errors by 90%"
    ),
    ("T. Rowe Price", 1): (
        "Redesigned legacy application with ~60% performance improvement and "
        "eliminated 3 recurring production incidents per month by implementing "
        "event-driven architecture on {cloud} using {services} and RDS-backed "
        "services"
    ),
    ("T. Rowe Price", 2): (
        "Rebuilt 4 critical {lang1} services using {lang2} to enable "
        "disaster-recovery failover, reducing potential downtime from 8 hours "
        "to under 30 minutes with automated PostgreSQL backup and failover "
        "systems"
    ),
    ("T. Rowe Price", 3): (
        "Developed high-performance data loaders integrating Active Directory "
        "data into {database} and {search}, optimising search functionality "
        "with JSON-based APIs and reducing average query response times by "
        "65% for 10,000+ daily users"
    ),
    ("Amazon Web Services", 0): (
        "Optimised AWS region build and Service Catalog deployment pipelines, "
        "reducing provisioning time by 40-55% across 15+ services while "
        "implementing automated validation and security controls"
    ),
    ("Amazon Web Services", 1): (
        "Deployed Service Catalog services across 5 newly launched AWS "
        "Regions (UAE, Melbourne, Spain, Zurich, Hyderabad) supporting "
        "global expansion"
    ),
    ("Amazon Web Services", 2): (
        "Led security escalation response managing 2,400+ hosts, implementing "
        "automated security patching pipelines"
    ),
}


class BatchCVGenerator:
    """Handles batch generation of CVs and cover letters using Claude 3.5 Sonnet"""
    
    def __init__(self, api_key: str):
        self.client = anthropic.Anthropic(api_key=api_key)
        self.skill_tiers = get_skills_by_tier()
        self.approved_skills = get_approved_skills()
        self.blacklisted_skills = get_blacklisted_skills()
        self.experiences = get_experiences()
        self.settings = get_settings()
    
    def parse_job_description(self, job_text: str) -> dict:
        """Parse job description using Claude"""
        prompt = f"""
        Analyze this job description and extract key information in JSON format.
        
        Job Description:
        {job_text}

        Return ONLY a valid JSON object with these fields:
        - job_title: The job title/position
        - company_name: Company name (if mentioned, otherwise "Unknown Company")
        - location: Location/city
        - required_skills: Array of technical skills mentioned as required
        - preferred_skills: Array of nice-to-have skills
        - years_experience: Experience level required
        - key_responsibilities: Array of main job responsibilities (3-5)
        - industry: Industry sector
        - remote_policy: "Remote", "Hybrid", "On-site", or "Not specified"
        """
        
        response = self.client.messages.create(
            model=CLAUDE_MODEL_QUALITY,
            max_tokens=1500,
            messages=[{"role": "user", "content": prompt}]
        )
        
        response_text = response.content[0].text
        json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
        if json_match:
            return json.loads(json_match.group())
        raise ValueError("Could not parse job description")
    
    def _filter_and_match_skills(self, job_skills: list) -> list:
        """
        Filter job skills to only include approved ones, prioritized by tier.
        This is the smart matching - if job mentions Kafka/SQS/etc and we have it approved, include it.
        """
        matched_skills = {
            'tier_1_core': [],
            'tier_2_major': [],
            'tier_3_specialist': [],
            'tier_4_tools': [],
        }
        
        for job_skill in job_skills:
            job_skill_lower = job_skill.lower().strip()
            
            # Check against each tier
            for tier, tier_skills in self.skill_tiers.items():
                for my_skill in tier_skills:
                    my_skill_lower = my_skill.lower()
                    
                    # Flexible matching
                    if (job_skill_lower in my_skill_lower or 
                        my_skill_lower in job_skill_lower or
                        job_skill_lower == my_skill_lower):
                        
                        # Check not blacklisted
                        if my_skill not in self.blacklisted_skills:
                            if my_skill not in matched_skills[tier]:
                                matched_skills[tier].append(my_skill)
                        break
        
        # Flatten in priority order
        result = []
        for tier in ['tier_1_core', 'tier_2_major', 'tier_3_specialist', 'tier_4_tools']:
            result.extend(matched_skills[tier])
        
        return result
    
    def _build_experience_context(self, matched_skills: list) -> str:
        """Build experience context with mutable tech stacks based on matched skills"""
        context_parts = []
        
        for exp in self.experiences:
            exp_context = f"\n{exp['company']} - {exp['role']}"
            if exp['start_date']:
                exp_context += f" ({exp['start_date']} - {'Present' if exp['is_current'] else exp['end_date']})"
            
            bullets = []
            for bp_idx, bp in enumerate(exp.get('bullet_points', [])):
                override_key = (exp['company'], bp_idx)
                desc = BULLET_OVERRIDES.get(override_key, bp['base_description'])
                
                # If there are tech placeholders, try to fill them with matched skills
                if bp.get('tech_placeholder'):
                    placeholders = bp['tech_placeholder'].split(',')
                    for ph in placeholders:
                        ph = ph.strip()
                        # Find a relevant skill from matched skills
                        replacement = self._find_skill_for_placeholder(ph, matched_skills)
                        if replacement:
                            desc = desc.replace('{' + ph + '}', replacement)
                        else:
                            # Use a generic term if no match
                            desc = desc.replace('{' + ph + '}', ph.title())
                
                bullets.append(f"  - {desc}")
            
            exp_context += "\n" + "\n".join(bullets)
            context_parts.append(exp_context)
        
        return "\n".join(context_parts)
    
    def _find_skill_for_placeholder(self, placeholder: str, matched_skills: list) -> str:
        """Find a matching skill for a placeholder based on context"""
        placeholder_lower = placeholder.lower()
        
        # Mapping of placeholder types to skill categories
        mappings = {
            'tech': ['Python', 'Java', 'JavaScript'],
            'lang1': ['Java', 'Python'],
            'lang2': ['Python', 'Java'],
            'cloud': ['AWS', 'Lambda', 'EC2'],
            'services': ['Lambda', 'SQS', 'S3', 'RDS'],
            'database': ['RDS', 'PostgreSQL', 'MySQL', 'DynamoDB'],
            'search': ['OpenSearch', 'Elasticsearch'],
            'messaging': ['SQS', 'Kafka', 'RabbitMQ'],
        }
        
        # Check if placeholder matches a mapping
        if placeholder_lower in mappings:
            for skill in mappings[placeholder_lower]:
                if skill in matched_skills:
                    return skill
        
        # Otherwise try direct match
        for skill in matched_skills:
            if placeholder_lower in skill.lower():
                return skill
        
        return None
    
    def generate_cv_and_cover_letter(self, job_info: dict, variant: str = 'professional') -> dict:
        """Generate both CV content and cover letter in a single LLM call"""
        
        # Get matched skills from job requirements
        all_job_skills = job_info.get('required_skills', []) + job_info.get('preferred_skills', [])
        matched_skills = self._filter_and_match_skills(all_job_skills)
        
        # Build experience context with matched tech
        experience_context = self._build_experience_context(matched_skills)
        
        # Get settings
        settings = self.settings
        expertise_count = int(settings.get('expertise_count', '14'))
        
        # Get variant tone instruction
        variant_info = CV_VARIANTS.get(variant, CV_VARIANTS['professional'])
        tone_instruction = variant_info['tone_instruction']
        
        skills_section = f"""
        MY APPROVED SKILLS (by priority tier):
        Tier 1 Core: {', '.join(self.skill_tiers.get('tier_1_core', []))}
        Tier 2 Major: {', '.join(self.skill_tiers.get('tier_2_major', []))}
        Tier 3 Specialist: {', '.join(self.skill_tiers.get('tier_3_specialist', []))}
        Tier 4 Tools: {', '.join(self.skill_tiers.get('tier_4_tools', []))}
        
        BLACKLISTED SKILLS (never mention these):
        {', '.join(self.blacklisted_skills[:30])}... (and more)
        
        MATCHED SKILLS FOR THIS JOB (prioritized):
        {', '.join(matched_skills)}
        """
        
        # Sanitize job title
        forbidden_keywords = r'\b(Senior|Lead|Principal|Staff|I|II|III|IV|DevOps)\b'
        sanitized_title = re.sub(forbidden_keywords, '', job_info.get('job_title', ''), flags=re.IGNORECASE).strip()
        
        # Get current experience info
        experiences = get_experiences()
        current_exp = next((e for e in experiences if e.get('is_current')), None)
        current_company = current_exp['company'] if current_exp else 'Compare the Market'
        
        # Extract ATS phrases from raw job text if available
        raw_job_text = job_info.get('raw_text', '')
        ats_phrases = []
        if raw_job_text:
            # Use the module-level function directly (no self-import needed)
            ats_phrases = extract_ats_phrases(raw_job_text)
        
        prompt = f"""
        Generate BOTH a customized CV and cover letter for this job application in a SINGLE response.
        
        WRITING STYLE: {tone_instruction}
        
        ATS REQUIREMENTS (be ATS-friendly without sacrificing readability):
        1. Include these phrases somewhere across the CV (use proper capitalisation - never lowercase tech terms): {', '.join(ats_phrases[:10]) if ats_phrases else ', '.join(job_info.get('required_skills', [])[:10])}
        2. Most bullets should include a quantified outcome (%, time, count) BUT architectural/leadership bullets may omit numbers when they capture clear scope/impact (mirror the model CV style)
        3. Use standard section headers only
        4. ALWAYS use proper capitalisation: Docker (not docker), Kubernetes (not kubernetes), CI/CD (not ci/cd), AWS, FastAPI, PostgreSQL, GitHub Actions, AI, LLM, etc.
        5. Mirror the exact language from the job description where natural

        TARGET JOB:
        - Position: {sanitized_title} at {job_info.get('company_name', 'Unknown Company')}
        - Location: {job_info.get('location', 'Not specified')}
        - Industry: {job_info.get('industry', 'Technology')}
        - Required Skills (MUST appear in CV): {', '.join(job_info.get('required_skills', []))}
        - Preferred Skills: {', '.join(job_info.get('preferred_skills', []))}
        - Key Responsibilities: {', '.join(job_info.get('key_responsibilities', []))}

        {skills_section}

        MY EXPERIENCE (adapt and enhance these):
        {experience_context}

        METRIC ANCHORS (these are REAL numbers from my work - use them EXACTLY, never round down):
        - Compare the Market: 50+ PRDs, 2 hours to 15 minutes, 6+ PMs, Head of Product, 7-agent pipeline, 34,000+ merge requests, ~95% adoption, 400+ engineers, 100+ engineers mentored, 5 workshops, 2 hackathons
        - T. Rowe Price: 90% migration error reduction, ~60% performance improvement, 3 incidents/month eliminated, 4 critical services rebuilt, 8 hours to under 30 minutes DR, 65% query time reduction, 10,000+ daily users
        - AWS: 40-55% provisioning time reduction, 15+ services, 5 AWS Regions, 2,400+ hosts

        MY BACKGROUND FOR COVER LETTER:
        - {settings.get('years_experience', '3.5')} years experience as Software Engineer (since July 2022)
        - Currently at {current_company} as Software Engineer - AI Native (since October 2025)
        - Previously at T. Rowe Price (financial services, enterprise data migration)
        - Previously at AWS (cloud infrastructure, systems automation)
        - Key focus: AI/ML automation, LangChain, LangGraph, Python, AWS
        - Education: {settings.get('education', 'BSc Cyber Security from Warwick University (2022)')}
        - Location: {settings.get('user_location', 'London, UK')}

        CRITICAL INSTRUCTIONS:
        1. CV must fit on exactly 1 page - bio HARD LIMIT 290 chars (count carefully), bullets 25-40 words each (model CVs average 28-32 words per bullet)
        2. Bio: 2-3 sentences. Sentence 1 = role + years + 2-3 domain areas (e.g., "Software Engineer with experience building Python-based backend services, enterprise platform tooling, and AI-driven automation across financial services and cloud infrastructure environments."). Sentence 2 = "Strong background in <3-5 CONCRETE TECHNOLOGIES from MATCHED SKILLS>, <one capability area>". NEVER use "Senior", "leader", "lead", or inflated titles. NEVER fill the bio with soft-skill JD phrases like "product ownership mindset", "adaptability", "growth mindset", "iterative development and prototyping" - those belong in the cover letter, not the CV bio.
        3. Bullets: short, crisp, action-verb led, 25-40 words each. Each bullet should pair a concrete artefact with 1-3 specific technologies and a clear outcome.
        4. Expertise: Exactly {expertise_count} CONCRETE TECHNOLOGIES drawn ONLY from MY APPROVED SKILLS / MATCHED SKILLS. Programming languages first. Group related items where natural (e.g. "Docker & Kubernetes", "AWS (ECS, Lambda, RDS, S3, SQS)"). DO NOT use soft skills, methodologies or JD phrasings like "Software Engineering", "Modern OO language proficiency", "Architectural design", "Iterative/Agile development", "Prototyping", "Cross-functional collaboration", "Scalable design patterns", "Product Development". Those belong in the bio/bullets, not expertise.
        5. Tech stacks: 8-11 CONCRETE TECHNOLOGIES (no soft skills, methodologies or process terms) from MATCHED SKILLS that appear in the job description, comma-separated. Always proper-cased.
        6. NEVER mention any BLACKLISTED SKILLS
        6a. NEVER claim languages/frameworks the candidate does not actually have. Only use technologies from MY APPROVED SKILLS. If the JD requires C#/.NET, Go, etc. and they are not in MY APPROVED SKILLS, DO NOT add them to the bio or expertise - just emphasise transferable strengths instead.
        7. STRICTLY AVOID filler phrases that don't add concrete information: "scalable design patterns", "distributed systems architecture", "supporting payment systems knowledge gained", "within scalable design patterns", "across enterprise financial services workflows". Be concrete instead.
        8. NEVER pad verbs: write "Led architecture of..." not "Led software engineering architecture of...". Write "Built..." not "Architected end-to-end design and delivery of...".
        9. Cover letter: 2-3 paragraphs, professional, mention 2-3 relevant matched skills.

        ==== ANTI-AI-CV STYLE RULES (CRITICAL - 2026 recruiters can spot generated CVs in seconds) ====

        THE CORE PRINCIPLE: it is not WHICH words you use, it is the DENSITY of buzzwords per sentence and whether each fancy word EARNS ITS PLACE next to a concrete implementation detail.

        A real senior engineer might write "Architected an event-driven migration pipeline on AWS using SQS and Lambda" - that is fine because there is a specific system, named tech, and clear architecture context.
        An AI writes "Architected scalable multi-agent orchestration systems enabling intelligent automation workflows across enterprise domains" - 4 buzzwords stacked, zero concrete content. This is what we are eliminating.

        THE STRIPPED-SENTENCE TEST (apply this to every bullet before writing it):
        Strip out adjectives and buzzwords (scalable, intelligent, advanced, robust, seamless, production-grade, agentic, AI-native, cutting-edge, next-generation). If the remaining sentence still describes something concrete and impressive, the bullet passes. If almost nothing is left, the bullet is buzzword-stacking - rewrite it.
          PASSES: "Built Python services handling automated PR review across 300+ engineers." -> stripped: "Built services handling PR review across 300+ engineers." Still strong.
          FAILS:  "Built scalable AI-native orchestration systems."                              -> stripped: "Built systems." Nothing left.

        A. HARD BANS - these words have effectively zero legitimate use in a CV and immediately read as LLM output. NEVER write them:
           leveraged / leveraging
           intelligent decomposition, intelligent routing, intelligent task routing, intelligent orchestration
           cutting-edge, next-generation, transformational, groundbreaking, state-of-the-art
           innovative, advanced AI, holistically, synergies, robust solutions
           seamlessly (the adverb; "seamless" the adjective is allowed once in context)

        B. DENSITY-CONTROLLED WORDS - legitimate when EARNED, dangerous when stacked. Allowed but each bullet may use AT MOST ONE of these, and ONLY if the same bullet also contains a concrete anchor (a specific named system, named tech, or a real number):
           architected / orchestrated / scalable / agentic / AI-native / seamless / production-grade / production-ready / distributed / intelligent / end-to-end / cross-functional / cross-team

        C. BUZZWORD-STACKING CHECK - across the whole CV, NEVER produce a single bullet that contains 2+ of the words in list B. Example of what to NEVER write: "scalable agentic orchestration", "intelligent end-to-end automation", "production-grade scalable AI-native pipeline". If you find yourself writing one of these phrases, replace it with a single concrete fact (the system name, the user count, the actual technology).

        D. VERB VARIETY - do NOT start more than 2 bullets with the same verb. Mix from:
           Built, Developed, Redesigned, Improved, Migrated, Integrated, Deployed, Reduced,
           Automated, Refactored, Led, Designed, Shipped, Maintained, Owned, Rolled out,
           Mentored, Investigated, Rebuilt, Consolidated.
           "Led architecture of" and "Designed and shipped" are also fine (used by your model CVs).

        E. STRUCTURE VARIATION - at least 3 of the 9 bullets must lead with the OUTCOME or scale, not the verb:
           * "Reduced disaster-recovery downtime from 8 hours to under 30 minutes by rebuilding 4 critical Python services on FastAPI."
           * "Across 400+ engineers, deployed an AI code review system processing 34,000+ MRs at ~95% adoption."
           * "From 2 hours to 15 minutes per PRD: automated the PRD-to-JIRA pipeline using LangGraph and FastAPI, now used by 6+ PMs including Head of Product."

        F. METRIC REALISM - recruiters distrust suspiciously clean stats:
           PREFER: ranges ("40-55%"), time conversions ("from 2 hours to 15 minutes"), concrete counts ("50+ PRDs", "400+ engineers", "15+ services", "2,400+ hosts").
           For percentages: prefix estimates with "~" ("~95% adoption"), AND pair with a denominator/scale ("~95% adoption across 400+ engineers" - never bare "95%").
           At most 2 of the 9 bullets should be a clean single percentage; the rest must use ranges, counts, or time conversions.
           NEVER use orphan stats like "by 80%+" without scale - it reads as fabricated.
           NEVER use the same metric pattern in adjacent bullets ("by 75%" followed by "by 80%+" is a tell).

        G. TECH-MENTION CAP - the tech stack line already lists technologies; repeating them in every bullet is keyword spam:
           Each technology appears in AT MOST 2 bullets across the whole CV.
           Spine exceptions: "Python" up to 4 bullets, "AWS" up to 4 bullets, "FastAPI" up to 3 bullets.
           If a tech is in the section's "Key tech stack" line, mention it ONCE in that section's bullets where it matters most, not in every bullet.

        H. OPERATIONAL LANGUAGE - senior engineers OWN systems, they don't just BUILD them. At least 1-2 bullets should reflect ownership/operation:
           reliability, rollout, monitoring, incident reduction, developer workflows, on-call rotation, runbooks, deployment pipeline, operational tooling, production support, post-incident review, supportability, observability.
           This is especially important for platform/finance/infrastructure roles.

        I. JD TONE - calibrate vocabulary to the target company:
           AI-native / agent / orchestration vocabulary is fine for: Anthropic, OpenAI, Perplexity, AI-first startups.
           Restrained, operational vocabulary performs better for: Man Group, Bloomberg, JPMorgan, Thought Machine, enterprise fintech, banks, asset managers, hedge funds.
           Look at the target company "{job_info.get('company_name', '')}" and industry "{job_info.get('industry', '')}". If it is enterprise-fintech / bank / asset manager: bias toward "Built / Developed / Migrated / Maintained / Owned / Reduced" and concrete systems language. If it is AI startup: "Designed and shipped / Deployed / Led architecture of" with light AI vocabulary is fine.

        ==== END ANTI-AI-CV STYLE RULES ====

        STYLE ANCHOR - bullets in your output should sound like these (these are real model exemplars - mirror their crispness, specificity, verb variety and metric realism):
          * "Led architecture of cross-team AI automation converting 50+ Product Requirement Documents (PRDs) into JIRA-ready tickets using LangGraph and FastAPI, reducing planning time from 2 hours to 15 minutes - adopted by 6+ product managers including Head of Product."  (verb-first, time-conversion, scale denominator)
          * "Designed and shipped a 7-agent LangGraph-based AI pipeline with Redis and parallel Python workers, enabling document processing with automated task routing and a production-ready GitLab knowledge graph indexed on AWS EFS."  (no metric, architecture-focused)
          * "Deployed AI-powered code review system processing 34,000+ merge requests at ~95% adoption across 400+ engineers, auto-approving high-quality MRs and integrating internal APIs and knowledge-graph tooling for codebase-aware suggestions."  (scale denominator + ~ prefix)
          * "Redesigned legacy application with ~60% performance improvement and eliminated 3 recurring production incidents per month by implementing event-driven architecture on AWS SQS with RDS-backed services."  (operational outcome - incidents per month)
          * "Rebuilt 4 critical Python services on FastAPI to enable disaster-recovery failover, reducing potential downtime from 8 hours to under 30 minutes with automated PostgreSQL backup."  (outcome-first time conversion)
          * "Optimised AWS region build and Service Catalog deployment pipelines, reducing provisioning time by 40-55% across 15+ services while implementing automated validation and security controls."  (range, not clean %)
          * "Mentored 100+ engineers on AI-native development through 5 workshops and 2 hackathons, supporting teams shipping production features with AI-powered Python backends."  (concrete counts)

        TECH SUBSTITUTION for Compare the Market bullets - pick concrete technologies from MATCHED SKILLS, never leave placeholders in output:
           - AI framework -> LangGraph, LangChain (default: LangGraph)
           - Containerisation/orchestration -> Docker, Kubernetes, AWS ECS (default: Kubernetes)
           - Database -> PostgreSQL, AWS RDS (default: PostgreSQL)
           - Cache -> Redis (default: Redis)
           - AI tools -> Python, Claude API, LLM APIs (default: Python and LLM APIs)

        Return ONLY a JSON object with this exact structure:
        {{
            "cv": {{
                "bio": "Updated bio paragraph - 2-3 sentences, max 290 chars, ATS keywords woven in naturally",
                "expertise": ["List of exactly {expertise_count} skills from MATCHED SKILLS"],
                "c": {{
                    "skills": "Python, LangGraph, LangChain, AWS, Docker, FastAPI, PostgreSQL, Redis, JSON, APIs",
                    "bp1": "<25-40 word bullet for Compare The Market - first theme - matching the STYLE ANCHOR examples above. Concrete tech, no placeholders, no annotation text.>",
                    "bp2": "<25-40 word bullet for Compare The Market - first theme continued.>",
                    "bp3": "<25-40 word bullet for Compare The Market - SDLC theme.>",
                    "bp4": "<25-40 word bullet for Compare The Market - mentoring theme.>"
                }},
                "t": {{
                    "skills": "Comma-separated tech stack string using MATCHED SKILLS",
                    "bp1": "First bullet point with specific tech and metrics (max 300 chars)",
                    "bp2": "Second bullet point with specific tech and metrics (max 300 chars)",
                    "bp3": "Third bullet point with specific tech and metrics (max 300 chars)",
                    "bp4": "Fourth bullet point with specific tech and metrics (max 300 chars)"
                }},
                "a": {{
                    "skills": "Comma-separated tech stack string using MATCHED SKILLS",
                    "bp1": "Single concise impact-driven bullet combining infrastructure automation & global systems (max 300 chars)"
                }}
            }},
            "cover_letter": "Complete cover letter text (2-3 paragraphs). Use name: {settings.get('user_name', 'Drew Gillies')}. No placeholders."
        }}
        """
        
        def _call_llm(p: str) -> dict:
            resp = self.client.messages.create(
                model=CLAUDE_MODEL_QUALITY,
                max_tokens=3500,
                messages=[{"role": "user", "content": p}]
            )
            txt = resp.content[0].text
            txt = re.sub(r',\s*([\}\]])', r'\1', txt)
            m = re.search(r'\{.*\}', txt, re.DOTALL)
            if not m:
                raise ValueError("Could not parse CV generation response")
            return json.loads(m.group())

        result = _call_llm(prompt)

        # Anti-AI-CV validation + targeted retry (up to 1 retry)
        if 'cv' in result and isinstance(result['cv'], dict):
            issues = detect_hype_issues(result['cv'])
            hard_ban = issues.get('hard_ban_hits', [])
            stacked = issues.get('stacked_bullets', [])
            orphan = issues.get('orphan_pcts', [])
            over_verbs = issues.get('over_repeated_verbs', {})
            outcome_first = issues.get('outcome_first_count', 0)

            needs_retry = (
                len(hard_ban) > 0
                or len(stacked) > 0
                or len(orphan) > 2
                or any(c > 2 for c in over_verbs.values())
                or outcome_first < 2
            )

            if needs_retry:
                feedback_parts = ["The previous CV failed the anti-AI-CV checks. Regenerate the SAME JSON structure but fix these specific issues:"]
                if hard_ban:
                    feedback_parts.append(f"- HARD-BANNED words appeared (remove every instance): {sorted(set(hard_ban))}")
                if stacked:
                    feedback_parts.append(
                        "- Buzzword-stacked bullets (2+ density words in a single sentence - REWRITE each with a single concrete fact and named tech/number):\n  "
                        + "\n  ".join('"' + s['text'] + '"' for s in stacked)
                    )
                if len(orphan) > 2:
                    feedback_parts.append(
                        f"- Too many clean orphan percentages without a denominator ({orphan}). Replace most with ranges (40-55%), time conversions (from 2 hours to 15 minutes), or counts (50+ PRDs)."
                    )
                if any(c > 2 for c in over_verbs.values()):
                    repeats = {v: c for v, c in over_verbs.items() if c > 2}
                    feedback_parts.append(f"- Verbs starting too many bullets (max 2 each): {repeats}. Pick from: Built, Developed, Redesigned, Improved, Migrated, Integrated, Deployed, Reduced, Automated, Refactored, Designed, Shipped, Maintained, Owned, Rebuilt.")
                if outcome_first < 2:
                    feedback_parts.append("- Only {} bullets lead with an OUTCOME/scale. Rewrite at least 3 to start with the outcome (e.g. 'Reduced X from 8 hours to 30 minutes by ...', 'Across 400+ engineers, deployed ...').".format(outcome_first))

                retry_prompt = prompt + "\n\n==== RETRY FEEDBACK ====\n" + "\n".join(feedback_parts) + "\n\nReturn ONLY the corrected JSON, same structure as before."
                try:
                    retry_result = _call_llm(retry_prompt)
                    if 'cv' in retry_result and isinstance(retry_result['cv'], dict):
                        retry_issues = detect_hype_issues(retry_result['cv'])
                        # Accept the retry if it has strictly fewer hard bans + stacked bullets
                        before_score = len(hard_ban) + len(stacked) + max(0, len(orphan) - 2)
                        after_score = (
                            len(retry_issues.get('hard_ban_hits', []))
                            + len(retry_issues.get('stacked_bullets', []))
                            + max(0, len(retry_issues.get('orphan_pcts', [])) - 2)
                        )
                        if after_score <= before_score:
                            result = retry_result
                except Exception as e:
                    print(f"  ⚠ Retry failed, keeping original CV: {e}")

        # Apply tech-casing safety net on the (possibly retried) CV
        if 'cv' in result and isinstance(result['cv'], dict):
            result['cv'] = normalise_cv_casing(result['cv'])

        # Validate page length
        is_valid, pages, suggestions = validate_single_page(result.get('cv', {}))
        result['page_validation'] = {
            'is_valid': is_valid,
            'estimated_pages': pages,
            'suggestions': suggestions
        }

        return result
    
    def generate_all_variants(self, job_info: dict) -> dict:
        """Generate all CV variants for a job"""
        variants = {}
        for variant_key in CV_VARIANTS.keys():
            try:
                variants[variant_key] = self.generate_cv_and_cover_letter(job_info, variant_key)
            except Exception as e:
                variants[variant_key] = {'error': str(e)}
        return variants
    
    def create_cover_letter_pdf(self, cover_letter_text: str, job_info: dict, output_path: str):
        """Create a professional PDF cover letter"""
        settings = self.settings
        
        doc = SimpleDocTemplate(output_path, pagesize=letter,
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)
        
        styles = getSampleStyleSheet()
        header_style = ParagraphStyle(
            'CustomHeader',
            parent=styles['Normal'],
            fontSize=12,
            fontName='Helvetica-Bold',
            spaceAfter=12,
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            fontName='Helvetica',
            spaceAfter=12,
            leading=14,
        )
        
        story = []
        
        # Header
        header_text = f"""
        <b>{settings.get('user_name', 'Drew Gillies')}</b><br/>
        Software Engineer<br/>
        {settings.get('user_location', 'London, UK')}<br/>
        {settings.get('user_email', 'drew.gillies@hotmail.co.uk')}<br/>
        {settings.get('user_phone', '07950 298726')}<br/>
        {settings.get('user_linkedin', 'linkedin.com/in/drew-gillies')}
        """
        story.append(Paragraph(header_text, header_style))
        story.append(Spacer(1, 20))
        
        # Date
        date_text = datetime.now().strftime("%B %d, %Y")
        story.append(Paragraph(date_text, normal_style))
        story.append(Spacer(1, 12))
        
        # Company address
        company_name = job_info.get('company_name', 'Unknown Company')
        location = job_info.get('location', '')
        if company_name != "Unknown Company":
            address_text = f"""
            Hiring Manager<br/>
            {company_name}<br/>
            {location if location != "Not specified" else ""}
            """
            story.append(Paragraph(address_text, normal_style))
            story.append(Spacer(1, 12))
        
        # Subject
        job_title = job_info.get('job_title', 'Software Engineer')
        subject_text = f"<b>Re: {job_title} Position</b>"
        story.append(Paragraph(subject_text, normal_style))
        story.append(Spacer(1, 12))
        
        # Body
        paragraphs = cover_letter_text.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                clean_paragraph = paragraph.strip()
                clean_paragraph = clean_paragraph.replace('[Your name]', settings.get('user_name', 'Drew Gillies'))
                clean_paragraph = clean_paragraph.replace('[Your Name]', settings.get('user_name', 'Drew Gillies'))
                clean_paragraph = clean_paragraph.replace('Best regards,', '')
                clean_paragraph = clean_paragraph.replace('Sincerely,', '')
                
                if clean_paragraph:
                    story.append(Paragraph(clean_paragraph, normal_style))
                    story.append(Spacer(1, 12))
        
        story.append(Spacer(1, 12))
        doc.build(story)
    
    def create_cv_docx(self, cv_data: dict, job_info: dict, output_path: str, create_pdf: bool = True):
        """Create CV document from template with clean metadata"""
        if not Path(TEMPLATE_PATH).exists():
            raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")
        
        doc = Document(TEMPLATE_PATH)
        
        # Clean metadata - remove any AI/tool indicators
        core_props = doc.core_properties
        core_props.author = self.settings.get('user_name', 'Drew Gillies')
        core_props.title = f"CV - {self.settings.get('user_name', 'Drew Gillies')}"
        core_props.subject = job_info.get('job_title', 'Software Engineer')
        core_props.keywords = ''
        core_props.comments = ''
        core_props.category = ''
        core_props.last_modified_by = self.settings.get('user_name', 'Drew Gillies')
        
        # Prepare replacements
        replacements = {
            'bio': cv_data.get('bio', ''),
            'expertise': cv_data.get('expertise', []),
            # Compare the Market (current role)
            'c.skills': cv_data.get('c', {}).get('skills', ''),
            'c.bp1': cv_data.get('c', {}).get('bp1', ''),
            'c.bp2': cv_data.get('c', {}).get('bp2', ''),
            'c.bp3': cv_data.get('c', {}).get('bp3', ''),
            'c.bp4': cv_data.get('c', {}).get('bp4', ''),
            # T. Rowe Price
            't.skills': cv_data.get('t', {}).get('skills', ''),
            't.bp1': cv_data.get('t', {}).get('bp1', ''),
            't.bp2': cv_data.get('t', {}).get('bp2', ''),
            't.bp3': cv_data.get('t', {}).get('bp3', ''),
            't.bp4': cv_data.get('t', {}).get('bp4', ''),
            # AWS (single concise bullet for older experience)
            'a.skills': cv_data.get('a', {}).get('skills', ''),
            'a.bp1': cv_data.get('a', {}).get('bp1', ''),
        }
        
        # Split expertise for two columns
        expertise = cv_data.get('expertise', [])
        midpoint = (len(expertise) + 1) // 2
        replacements['expertise'] = expertise[:midpoint]
        replacements['expertise2'] = expertise[midpoint:]
        
        # Replace placeholders
        self._replace_placeholders(doc, replacements)
        
        doc.save(output_path)
        
        # Create PDF version
        if create_pdf:
            pdf_path = output_path.replace('.docx', '.pdf')
            self._convert_docx_to_pdf(output_path, pdf_path)
            return pdf_path
        return None
    
    def _convert_docx_to_pdf(self, docx_path: str, pdf_path: str):
        """Convert DOCX to PDF using available tools"""
        try:
            # Try using LibreOffice (most reliable cross-platform)
            result = subprocess.run([
                'soffice', '--headless', '--convert-to', 'pdf',
                '--outdir', str(Path(pdf_path).parent),
                docx_path
            ], capture_output=True, timeout=30)
            
            if result.returncode == 0:
                # LibreOffice names output based on input filename
                expected_pdf = Path(docx_path).with_suffix('.pdf')
                if expected_pdf.exists() and str(expected_pdf) != pdf_path:
                    expected_pdf.rename(pdf_path)
                return True
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
        
        try:
            # Try docx2pdf (requires MS Word on Mac/Windows)
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            return True
        except Exception:
            pass
        
        # If no PDF converter available, skip PDF creation
        return False
    
    def _replace_placeholders(self, doc, replacements):
        """Replace placeholders in document"""
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)
        
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                self._replace_in_paragraph(paragraph, replacements)
            for paragraph in section.footer.paragraphs:
                self._replace_in_paragraph(paragraph, replacements)
    
    def _replace_in_paragraph(self, paragraph, replacements):
        """Replace placeholders in a paragraph - handles split runs"""
        # First, try to find and replace in individual runs
        for key, value in replacements.items():
            placeholder = f"{{{{{key}}}}}"
            
            # Prepare replacement text
            if key in ['t.skills', 'a.skills']:
                if isinstance(value, list):
                    value_text = ', '.join(value)
                else:
                    value_text = str(value)
            elif isinstance(value, list):
                # Each skill on new line with bullet
                value_text = '\n• '.join(value)
                value_text = '• ' + value_text if value else ''
            else:
                value_text = str(value).replace('\n', ' ').strip()
            
            # Method 1: Direct run replacement
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value_text)
            
            # Method 2: If placeholder spans multiple runs, rebuild paragraph
            full_text = paragraph.text
            if placeholder in full_text:
                # Placeholder exists but wasn't in a single run - it's split
                new_text = full_text.replace(placeholder, value_text)
                # Clear all runs and set text on first run
                if paragraph.runs:
                    # Preserve formatting from first run
                    first_run = paragraph.runs[0]
                    for run in paragraph.runs[1:]:
                        run.text = ''
                    first_run.text = new_text


# URL Scraping functions

def scrape_job_url(url: str, use_llm: bool = True) -> str:
    """
    Scrape job description from a URL.
    If use_llm=True, uses Claude to extract and clean the job description.
    """
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
        }
        
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style", "nav", "header", "footer", "aside"]):
            script.decompose()
        
        # Try to find job-specific content containers
        job_content = None
        selectors = [
            '[class*="job-description"]',
            '[class*="jobDescription"]',
            '[class*="job_description"]',
            '[class*="job-details"]',
            '[class*="jobDetails"]',
            '[class*="posting-content"]',
            '[class*="description"]',
            '[id*="job-description"]',
            '[id*="jobDescription"]',
            'article',
            'main',
            '[role="main"]',
        ]
        
        for selector in selectors:
            elements = soup.select(selector)
            if elements:
                best = max(elements, key=lambda x: len(x.get_text()))
                if len(best.get_text().strip()) > 200:
                    job_content = best
                    break
        
        if job_content:
            text = job_content.get_text(separator='\n', strip=True)
        else:
            text = soup.get_text(separator='\n', strip=True)
        
        # Clean up
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        text = '\n'.join(lines)
        
        if len(text) > 12000:
            text = text[:12000] + '\n[Content truncated...]'
        
        # Use LLM to extract and clean the job description
        if use_llm and ANTHROPIC_API_KEY:
            text = llm_extract_job_description(text, url)
        
        return text
        
    except requests.exceptions.Timeout:
        raise Exception(f"Timeout while fetching {url}")
    except requests.exceptions.RequestException as e:
        raise Exception(f"Failed to fetch {url}: {str(e)}")
    except Exception as e:
        raise Exception(f"Error scraping {url}: {str(e)}")


def llm_extract_job_description(raw_text: str, url: str) -> str:
    """
    Use Claude to extract and clean job description from raw scraped text.
    This improves quality by removing navigation, ads, and irrelevant content.
    """
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    
    prompt = f"""
    Extract the job description from this scraped web page content.
    
    URL: {url}
    
    RAW CONTENT:
    {raw_text[:10000]}
    
    INSTRUCTIONS:
    1. Extract ONLY the job posting content - remove navigation, ads, footers, cookie notices
    2. Include: Job title, company name, location, requirements, responsibilities, benefits
    3. Preserve the original wording of requirements and skills
    4. Format cleanly with clear sections
    5. If salary/compensation is mentioned, include it
    6. Remove duplicate content
    
    Return ONLY the cleaned job description text, nothing else.
    """
    
    try:
        # Use Haiku for URL extraction - simple cleanup task
        response = client.messages.create(
            model=CLAUDE_MODEL_FAST,
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}]
        )
        
        extracted = response.content[0].text.strip()
        
        # Sanity check - if extraction is too short, return original
        if len(extracted) < 200:
            return raw_text
        
        return extracted
        
    except Exception as e:
        print(f"LLM extraction failed: {e}, using raw text")
        return raw_text


# Mapping of common lowercase tech tokens to their canonical capitalisation.
# Used as a safety net to ensure model bullets render with consistent casing
# (the ATS keyword list and required-skills list are lowercase, which can leak).
_TECH_CASE_MAP = {
    'docker': 'Docker', 'kubernetes': 'Kubernetes', 'k8s': 'Kubernetes',
    'ci/cd': 'CI/CD', 'ci cd': 'CI/CD', 'aws': 'AWS', 'gcp': 'GCP',
    'azure': 'Azure', 'fastapi': 'FastAPI', 'postgresql': 'PostgreSQL',
    'mysql': 'MySQL', 'mongodb': 'MongoDB', 'redis': 'Redis',
    'elasticsearch': 'Elasticsearch', 'opensearch': 'OpenSearch',
    'github actions': 'GitHub Actions', 'gitlab': 'GitLab',
    'rest api': 'REST API', 'restful api': 'RESTful API',
    'graphql': 'GraphQL', 'sql': 'SQL', 'json': 'JSON',
    'langgraph': 'LangGraph', 'langchain': 'LangChain',
    'jenkins': 'Jenkins', 'terraform': 'Terraform', 'lambda': 'Lambda',
    'cloudwatch': 'CloudWatch', 'sqs': 'SQS', 'sns': 'SNS', 's3': 'S3',
    'rds': 'RDS', 'ecs': 'ECS', 'ec2': 'EC2', 'efs': 'EFS',
    'pytest': 'pytest',  # intentionally lowercase
    'nodejs': 'Node.js', 'node.js': 'Node.js',
    'llm': 'LLM', 'llms': 'LLMs', 'api': 'API', 'apis': 'APIs',
    'mr': 'MR', 'mrs': 'MRs', 'prd': 'PRD', 'prds': 'PRDs',
}


def fix_tech_casing(text: str) -> str:
    """Replace lowercase tech tokens with their canonical capitalisation.
    Skips tokens that already appear inside an obviously-cased context (e.g. URLs)."""
    if not isinstance(text, str) or not text:
        return text
    # Sort by length desc so multi-word tokens replace first
    for token, canonical in sorted(_TECH_CASE_MAP.items(), key=lambda kv: -len(kv[0])):
        # Word-boundary, case-insensitive, but only replace when the match is not already canonical
        pattern = re.compile(r'(?<![A-Za-z0-9_/.-])' + re.escape(token) + r'(?![A-Za-z0-9_/.-])', re.IGNORECASE)
        def _sub(m):
            return canonical if m.group(0) != canonical else m.group(0)
        text = pattern.sub(_sub, text)
    return text


# ============================================================================
# Anti-AI-CV detection (informational - used by benchmark scorer)
# ============================================================================
# Two-tier model:
#   HARD_BANNED  = words with effectively zero legitimate use; any occurrence is a tell.
#   DENSITY_WORDS = legitimate when EARNED (paired with concrete anchor) but dangerous
#                   when stacked. The real signal is HOW MANY appear in one bullet, not
#                   that they appear at all.

HARD_BANNED = [
    r'leverag(?:e|ed|ing|es)',
    r'intelligent\s+(?:decomposition|routing|task\s+routing|orchestration)',
    r'innovative', r'cutting[-\s]edge', r'next[-\s]generation',
    r'transformational', r'groundbreaking', r'state[-\s]of[-\s]the[-\s]art',
    r'advanced\s+AI', r'holistically', r'synergies',
    r'robust\s+solutions?', r'seamlessly',
]

# Density-controlled (per-bullet). Allowed but stacking 2+ in one bullet = AI tell.
DENSITY_WORDS = [
    'architected', 'orchestrated', 'orchestration', 'scalable', 'agentic',
    'ai-native', 'seamless', 'production-grade', 'production-ready',
    'distributed', 'intelligent', 'end-to-end', 'cross-functional', 'cross-team',
]


def _bullet_buzzword_count(bullet: str) -> int:
    """Count DENSITY_WORDS occurrences in a single bullet (case-insensitive)."""
    b = bullet.lower()
    n = 0
    for w in DENSITY_WORDS:
        n += len(re.findall(r'(?<![A-Za-z])' + re.escape(w) + r'(?![A-Za-z])', b))
    return n


def _bullet_has_concrete_anchor(bullet: str) -> bool:
    """True if bullet contains a real number OR a recognisable named tech.
    Used by the stripped-sentence test - bullets without an anchor are pure buzzwords."""
    if re.search(r'\d', bullet):
        return True
    # Look for any canonical tech token (Docker, FastAPI, PostgreSQL, AWS, LangGraph, etc.)
    named_tech = ['Docker', 'Kubernetes', 'FastAPI', 'PostgreSQL', 'AWS', 'LangGraph',
                  'LangChain', 'Redis', 'Python', 'GitHub Actions', 'Jenkins', 'Lambda',
                  'CloudWatch', 'SQS', 'RDS', 'ECS', 'EC2', 'S3', 'Kafka', 'MySQL',
                  'MongoDB', 'Elasticsearch', 'Terraform', 'Flask', 'Django', 'Active Directory',
                  'JIRA', 'GitLab', 'CI/CD', 'PRD', 'PRDs', 'MR', 'MRs']
    for tech in named_tech:
        if re.search(r'(?<![A-Za-z0-9])' + re.escape(tech) + r'(?![A-Za-z0-9])', bullet):
            return True
    return False


def detect_hype_issues(cv: dict) -> dict:
    """Detect anti-AI-CV violations in a generated CV.
    Returns a dict with issue lists (informational - used by the benchmark scorer)."""
    bullets = []
    for sec in ('c', 't', 'a'):
        if isinstance(cv.get(sec), dict):
            for k, v in cv[sec].items():
                if k.startswith('bp') and isinstance(v, str):
                    bullets.append(v)
    bio = cv.get('bio', '') or ''
    all_text = ' '.join(bullets) + ' ' + bio

    # Hard bans (any occurrence anywhere)
    hard_ban_hits = []
    for pat in HARD_BANNED:
        for m in re.finditer(pat, all_text, re.IGNORECASE):
            hard_ban_hits.append(m.group(0))

    # Buzzword density per bullet (stacking detection)
    stacked_bullets = []
    bullet_buzz_counts = []
    for b in bullets:
        n = _bullet_buzzword_count(b)
        bullet_buzz_counts.append(n)
        if n >= 2:
            stacked_bullets.append({'count': n, 'text': b})

    # Stripped-sentence failures: bullets that contain density words but no concrete anchor
    no_anchor = []
    for b in bullets:
        if _bullet_buzzword_count(b) >= 1 and not _bullet_has_concrete_anchor(b):
            no_anchor.append(b)

    # Suspicious clean percentages: not prefixed with ~, not in a range, no nearby denominator
    orphan_pcts = []
    for b in bullets:
        for m in re.finditer(r'(?<!~)\b(\d{2,3})%\+?', b):
            start = m.start()
            # Skip ranges like 40-55%
            preceding = b[max(0, start - 4):start]
            if re.search(r'\d-$', preceding):
                continue
            # Skip if a denominator/scale (digits + noun) appears within 60 chars after
            tail = b[m.end():m.end() + 80]
            if re.search(r'\b\d[\d,]*\+?\s+\w+', tail):
                continue
            orphan_pcts.append(m.group(0))

    # Verb-start repetition (>2 bullets starting with the same verb)
    verb_starts = []
    for b in bullets:
        toks = b.strip().split()
        if toks:
            verb_starts.append(toks[0].rstrip(',').lower())
    verb_repeat = {}
    for v in verb_starts:
        verb_repeat[v] = verb_repeat.get(v, 0) + 1
    over_repeated_verbs = {v: c for v, c in verb_repeat.items() if c > 2}

    # Outcome-first bullets (not starting with a known verb stem)
    verb_set = {'built', 'developed', 'redesigned', 'improved', 'migrated', 'integrated',
                'deployed', 'reduced', 'automated', 'refactored', 'led', 'designed',
                'shipped', 'maintained', 'owned', 'rolled', 'mentored', 'investigated',
                'rebuilt', 'consolidated', 'architected', 'orchestrated', 'optimised',
                'optimized', 'delivered', 'implemented', 'established', 'created'}
    outcome_first_count = sum(1 for v in verb_starts if v not in verb_set)

    # Operational language presence
    op_terms = ['reliability', 'monitoring', 'incident', 'rollout', 'runbook',
                'on-call', 'on call', 'developer workflow', 'operational tooling',
                'production support', 'deployment pipeline', 'post-incident',
                'supportability', 'observability']
    op_hits = sum(1 for term in op_terms if term.lower() in all_text.lower())

    return {
        'hard_ban_hits': hard_ban_hits,
        'stacked_bullets': stacked_bullets,  # bullets with 2+ density words
        'bullet_buzz_counts': bullet_buzz_counts,
        'no_anchor_bullets': no_anchor,
        'orphan_pcts': orphan_pcts,
        'over_repeated_verbs': over_repeated_verbs,
        'outcome_first_count': outcome_first_count,
        'op_terms_hits': op_hits,
    }


def normalise_cv_casing(cv: dict) -> dict:
    """Apply tech-casing fix to all text fields in a CV dict."""
    if not isinstance(cv, dict):
        return cv
    if 'bio' in cv:
        cv['bio'] = fix_tech_casing(cv.get('bio', ''))
    if 'expertise' in cv and isinstance(cv['expertise'], list):
        cv['expertise'] = [fix_tech_casing(s) for s in cv['expertise']]
    for sec in ('c', 't', 'a'):
        if sec in cv and isinstance(cv[sec], dict):
            for k, v in list(cv[sec].items()):
                if isinstance(v, str):
                    cv[sec][k] = fix_tech_casing(v)
    return cv


def extract_ats_phrases(job_text: str) -> List[str]:
    """Extract high-value ATS phrases from job description"""
    job_lower = job_text.lower()
    phrases = []
    
    # Known tech phrases
    tech_phrases = [
        'aws lambda', 'aws s3', 'aws dynamodb', 'aws sqs', 'ci/cd', 'ci cd',
        'github actions', 'docker', 'kubernetes', 'terraform', 'postgresql',
        'restful api', 'rest api', 'microservices', 'data pipeline', 'etl',
        'agile', 'scrum', 'unit testing', 'integration testing', 'python',
        'java', 'javascript', 'typescript', 'react', 'node.js', 'sql',
        'cloud infrastructure', 'infrastructure as code', 'event-driven',
        'distributed systems', 'api design', 'system design', 'data migration',
    ]
    
    for phrase in tech_phrases:
        if phrase in job_lower:
            phrases.append(phrase)
    
    return phrases[:15]  # Top 15 phrases


def llm_enhance_cv_content(cv_data: dict, job_info: dict, variant: str) -> dict:
    """
    Use an additional LLM call to enhance and validate CV content quality.
    Ensures bullet points have metrics, bio is compelling, and skills are relevant.
    Uses exact phrases from job description for ATS optimization.
    """
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    
    required_skills = job_info.get('required_skills', [])
    preferred_skills = job_info.get('preferred_skills', [])
    responsibilities = job_info.get('key_responsibilities', [])
    
    # Extract exact phrases from job for ATS matching
    job_text = job_info.get('raw_text', '')
    ats_phrases = extract_ats_phrases(job_text) if job_text else required_skills[:10]
    
    prompt = f"""
    CRITICAL: Optimize this CV for ATS (Applicant Tracking Systems). 
    
    TARGET JOB:
    - Title: {job_info.get('job_title')}
    - Company: {job_info.get('company_name')}
    - Industry: {job_info.get('industry', 'Technology')}
    
    EXACT PHRASES TO INCLUDE (use these verbatim for ATS matching):
    {', '.join(ats_phrases)}
    
    REQUIRED SKILLS (must appear in CV):
    {', '.join(required_skills[:10])}
    
    KEY RESPONSIBILITIES FROM JOB (mirror this language):
    {'; '.join(responsibilities[:5]) if responsibilities else 'Not specified'}
    
    CURRENT CV CONTENT:
    {json.dumps(cv_data, indent=2)}
    
    OPTIMISATION RULES (preserve the existing crisp style; only fix gaps - DO NOT add skills/languages the candidate doesn't have):
    1. Most bullets should have a quantified outcome, but DO NOT force metrics into architectural/leadership bullets if it makes them awkward - mirror the model CV style.
    2. Ensure these phrases appear naturally somewhere in the CV (proper-cased), BUT ONLY if the candidate actually has them: {', '.join(ats_phrases[:12]) if ats_phrases else ', '.join(required_skills[:12])}
    3. Distribute required skills across bullets but keep each bullet focused on one theme.
    4. ALWAYS use proper capitalisation: Docker (not docker), Kubernetes (not kubernetes), CI/CD (not ci/cd), GitHub Actions, FastAPI, PostgreSQL, AWS, REST API, LLM.
    5. Use a MIX of verb stems (don't start more than 2 bullets with the same verb): Built, Developed, Redesigned, Improved, Migrated, Integrated, Deployed, Reduced, Automated, Refactored, Led, Designed, Shipped, Maintained, Owned, Rolled out, Mentored, Investigated, Rebuilt. Do NOT pad verbs ("Led software engineering architecture of" -> "Led architecture of").
    6. Each bullet 25-40 words. Vary the structure: at least 3 of 9 bullets must lead with the OUTCOME (e.g. "Reduced downtime from 8 hours to under 30 minutes by ...", "Across 400+ engineers, deployed ...") rather than a verb. Don't repeat the same Verb+Tech+Metric formula every bullet.
    7. STRICTLY remove filler phrases: "scalable design patterns", "distributed systems architecture", "supporting payment systems knowledge gained", "within scalable design patterns", "across enterprise financial services workflows", "product development cycles", "product ownership mindset", "growth mindset", "collaboration with cross-functional teams" (in bio/expertise), "iterative development and prototyping" (in bio/expertise). Replace with concrete content or shorten.
    7a. HARD BANS (effectively zero legitimate use - immediately read as LLM): "leveraged", "leveraging", "intelligent decomposition", "intelligent routing", "intelligent orchestration", "innovative", "cutting-edge", "next-generation", "transformational", "groundbreaking", "state-of-the-art", "advanced AI", "holistically", "synergies", "robust solutions", "seamlessly". Replace with simpler verbs and concrete nouns.
    7b. DENSITY CHECK (the core anti-AI-CV rule): these words are ALLOWED but only when EARNED. Each bullet may contain AT MOST ONE of them, AND only if the same bullet has a concrete anchor (named system, named tech, or real number). The words: "architected", "orchestrated", "scalable", "agentic", "AI-native", "seamless", "production-grade", "production-ready", "distributed", "intelligent", "end-to-end", "cross-functional", "cross-team". NEVER write a bullet that contains 2+ of these (e.g. "scalable AI-native orchestration" is forbidden - that is buzzword stacking).
    7b2. STRIPPED-SENTENCE TEST: for each bullet, mentally strip adjectives and the words above. If little concrete content remains, the bullet is buzzword-stacking - rewrite it with a specific system, named tech, or real number.
    7c. METRIC REALISM: PREFER ranges ("40-55%"), time conversions ("from 2 hours to 15 minutes"), concrete counts ("50+ PRDs", "400+ engineers", "15+ services"). For any percentage that is an estimate, prefix with "~" (e.g. "~95% adoption"). ALWAYS pair a percentage with a denominator/scale ("~95% adoption across 400+ engineers") - never bare "95%". Only ~2 of the 9 bullets should be a clean percentage; the rest should use ranges, counts, or time conversions. Reject orphan stats like "by 80%+" without scale.
    7d. TECH-MENTION CAP: each tech may appear in at most 2 bullets across the whole CV; spine techs "Python" (up to 4 bullets), "AWS" (up to 4), "FastAPI" (up to 3). Don't repeat techs already named in the section's "Key tech stack" line in every bullet of that section.
    7e. OPERATIONAL LANGUAGE: at least 1 bullet should mention reliability/monitoring/incident reduction/rollout/runbooks/developer workflows/on-call (especially for platform/finance/infra roles). Senior engineers OWN systems.
    8. BIO HARD CONSTRAINTS:
       - Maximum 290 characters total.
       - Sentence 1: role + years + 2-3 domain areas (financial services, cloud infrastructure, AI automation, etc.).
       - Sentence 2: "Strong background in <3-5 CONCRETE TECHNOLOGIES the candidate has>, <one capability area>".
       - NEVER use "Senior", "leader", "lead", "Expert in".
       - NEVER claim languages/frameworks NOT in the input CV (no fabricating C#/.NET, Go, Rust, etc.).
       - NEVER fill bio with soft-skill JD phrases ("product ownership mindset", "growth mindset", "adaptability").
    9. EXPERTISE HARD CONSTRAINTS:
       - Concrete technologies only (Python, AWS, FastAPI, Docker, Kubernetes, PostgreSQL, Redis, GitHub Actions, Terraform, etc.) and possibly capability domains seen in model CVs ("Event-Driven Architectures", "Platform Engineering", "CI/CD & DevOps Practices", "AI Systems & LLM Integrations", "APIs & Backend Services", "Infrastructure as Code (Terraform)").
       - NEVER include: "Software Engineering", "Architectural design", "Iterative/Agile development", "Prototyping", "Cross-functional collaboration", "Product Development", "Modern OO language proficiency", "Scalable design patterns", "Automated testing", or any language/framework not in the input CV.

    PHRASE PLACEMENT GUIDE (only use phrases the candidate actually has - skip required skills they lack):
    - Bio: weave in 2-3 of these where natural: {', '.join(required_skills[:3])}
    - T. Rowe Price bullets: weave in if applicable: {', '.join(required_skills[3:7]) if len(required_skills) > 3 else ''}
    - AWS bullet: weave in if applicable: {', '.join(required_skills[7:11]) if len(required_skills) > 7 else ''}
    - Expertise list: lead with concrete TECHNOLOGIES only (Python, AWS, FastAPI, Docker, Kubernetes, PostgreSQL, etc.) - NOT soft skills, methodologies or JD phrasings like "Software Engineering", "Architectural design", "Iterative/Agile development", "Prototyping", "Cross-functional collaboration".

    STYLE ANCHOR (mirror this crisp tone):
    - "Led architecture of cross-team AI automation converting 50+ Product Requirement Documents (PRDs) into JIRA-ready tickets using LangGraph and FastAPI, reducing planning time from 2 hours to 15 minutes - adopted by 6+ product managers including Head of Product."
    - "Designed and shipped a 7-agent LangGraph-based AI pipeline with Redis and parallel Python workers, enabling end-to-end PRD processing with automated task routing and production-ready GitLab knowledge graph indexing on AWS EFS."
    - "Redesigned legacy application with 60% performance improvement and eliminated 3 recurring production incidents per month by implementing scalable event-driven architecture on AWS using SQS and RDS-backed services."
    - "Optimised AWS region build and Service Catalog deployment pipelines, reducing provisioning time by 40-55% across 15+ services while implementing automated validation and security controls."
    
    Return enhanced CV as JSON:
    {{
        "bio": "ATS-optimized bio mentioning key technologies",
        "expertise": ["skill1", "skill2", ...11 skills],
        "c": {{"skills": "LangChain, LangGraph, Python, AWS, Redis, PostgreSQL", "bp1": "...", "bp2": "...", "bp3": "...", "bp4": "..."}},
        "t": {{"skills": "comma-separated tech stack", "bp1": "...", "bp2": "...", "bp3": "...", "bp4": "..."}},
        "a": {{"skills": "comma-separated tech stack", "bp1": "Single concise impact-driven bullet"}}
    }}
    
    Return ONLY valid JSON.
    """
    
    try:
        # Use Haiku for enhancement - structured improvement task (called 3x per job)
        response = client.messages.create(
            model=CLAUDE_MODEL_FAST,
            max_tokens=2500,
            messages=[{"role": "user", "content": prompt}]
        )
        
        response_text = response.content[0].text
        # Clean and parse JSON
        response_text = re.sub(r',\s*([\}\]])', r'\1', response_text)
        json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
        
        if json_match:
            enhanced = json.loads(json_match.group())
            enhanced = normalise_cv_casing(enhanced)
            # Safety net: only accept the enhanced CV if it does NOT have more
            # anti-AI-CV issues than the pre-enhancement CV. Otherwise revert.
            try:
                before = detect_hype_issues(cv_data)
                after = detect_hype_issues(enhanced)
                before_score = (
                    len(before.get('hard_ban_hits', [])) * 3
                    + len(before.get('stacked_bullets', [])) * 3
                    + len(before.get('no_anchor_bullets', []))
                    + max(0, len(before.get('orphan_pcts', [])) - 2)
                )
                after_score = (
                    len(after.get('hard_ban_hits', [])) * 3
                    + len(after.get('stacked_bullets', [])) * 3
                    + len(after.get('no_anchor_bullets', []))
                    + max(0, len(after.get('orphan_pcts', [])) - 2)
                )
                if after_score > before_score:
                    print(f"    ⚠ Enhancement degraded style (issues {before_score} -> {after_score}), reverting")
                    return cv_data
            except Exception:
                pass
            print(f"    ✓ CV content enhanced (Haiku)")
            return enhanced

        return cv_data
        
    except Exception as e:
        print(f"    ⚠ Enhancement failed: {e}, using original")
        return cv_data


def llm_judge_cv_feedback(cv_data: dict, job_info: dict) -> dict:
    """
    LLM-as-judge: Opus reviews the final CV against the JD and returns a
    corrected CV with specific fixes. Only modifies bullets/bio/expertise that
    need improvement — preserves everything else.
    """
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

    job_title = job_info.get('job_title', '')
    company = job_info.get('company_name', '')
    required = job_info.get('required_skills', [])
    preferred = job_info.get('preferred_skills', [])
    responsibilities = job_info.get('key_responsibilities', [])
    raw_jd = job_info.get('raw_text', '')

    prompt = f"""You are an expert CV reviewer. You will receive a generated CV (as JSON) and the target job description.

Your job: review the CV and return a CORRECTED version. Fix only what needs fixing. Preserve what is already good.

TARGET JOB:
- Position: {job_title} at {company}
- Required Skills: {', '.join(required)}
- Preferred Skills: {', '.join(preferred)}
- Responsibilities: {', '.join(responsibilities[:5])}

FULL JOB DESCRIPTION:
{raw_jd[:3000]}

CURRENT CV:
{json.dumps(cv_data, indent=2)}

REVIEW CHECKLIST - fix any of these you find:
1. METRIC ACCURACY: if a bullet says "100+ engineers" or "34,000+ merge requests" or "10,000+ daily users" or "65%", keep those exact numbers. Never deflate (e.g. turning 10,000+ into 2,000+ or 34,000+ into 30,000+).
2. EXPERTISE MUST BE CONCRETE TECHNOLOGIES ONLY: every item must be something you can install/import (Python, FastAPI, Docker, LangGraph, etc). Remove any soft skills, methodologies or JD phrasings like "System Design and Architecture", "Testing and Clean Code Practices", "Full-Stack Development", "Event-Driven Architecture". Replace with concrete tech from the CV's bullet points.
3. BULLET ATTRIBUTION: Compare the Market bullets (c section) should only contain Compare the Market work. T. Rowe Price bullets (t section) should only contain T. Rowe Price work. AWS bullets (a section) should only contain AWS work. No cross-contamination.
4. MISSING JD KEYWORDS: if important required skills from the JD are missing from the CV and the candidate plausibly has them, weave them in naturally.
5. BIO: max 290 chars, 2-3 sentences, concrete technologies, no inflated titles.
6. EACH BULLET: 25-40 words, action-verb led, with specific tech and outcome.
7. NO BUZZWORD STACKING: never 2+ of these in one bullet: architected, orchestrated, scalable, agentic, AI-native, production-grade, distributed, end-to-end, cross-functional.
8. NO HARD-BANNED WORDS: leveraged, leveraging, cutting-edge, next-generation, innovative, groundbreaking, state-of-the-art, holistically, synergies, robust solutions, seamlessly.

Return the corrected CV as JSON with the EXACT same structure as the input. Change ONLY what needs fixing. If the CV is already good, return it unchanged.
Return ONLY valid JSON, no commentary."""

    try:
        response = client.messages.create(
            model=CLAUDE_MODEL_JUDGE,
            max_tokens=3500,
            messages=[{"role": "user", "content": prompt}]
        )

        txt = response.content[0].text
        txt = re.sub(r',\s*([\}\]])', r'\1', txt)
        json_match = re.search(r'\{.*\}', txt, re.DOTALL)

        if json_match:
            judged = json.loads(json_match.group())
            judged = normalise_cv_casing(judged)

            # Safety: verify the judged CV still has the expected sections
            for key in ('bio', 'expertise', 'c', 't', 'a'):
                if key not in judged:
                    print(f"    ⚠ Judge output missing '{key}', keeping original")
                    return cv_data

            # Safety: check anti-AI-CV issues didn't get worse
            try:
                before = detect_hype_issues(cv_data)
                after = detect_hype_issues(judged)
                before_score = (
                    len(before.get('hard_ban_hits', [])) * 3
                    + len(before.get('stacked_bullets', [])) * 3
                    + len(before.get('no_anchor_bullets', []))
                )
                after_score = (
                    len(after.get('hard_ban_hits', [])) * 3
                    + len(after.get('stacked_bullets', [])) * 3
                    + len(after.get('no_anchor_bullets', []))
                )
                if after_score > before_score:
                    print(f"    ⚠ Judge degraded style (issues {before_score} -> {after_score}), keeping original")
                    return cv_data
            except Exception:
                pass

            print(f"    ✓ CV refined by judge (Opus)")
            return judged

        return cv_data

    except Exception as e:
        print(f"    ⚠ Judge failed: {e}, using original")
        return cv_data


def llm_select_best_variant(variant_results: list, job_text: str, job_info: dict) -> dict:
    """
    Use LLM to analyze all CV variants and select the best one for the job.
    Returns dict with recommended variant and reasoning.
    """
    if not variant_results or len(variant_results) < 2:
        # Only one variant, no need to compare
        return {
            'recommended': variant_results[0]['variant'] if variant_results else 'professional',
            'reasoning': 'Only one variant generated.',
            'confidence': 'high'
        }
    
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    
    # Build variant summaries for comparison
    variants_text = ""
    for vr in variant_results:
        if vr.get('error'):
            continue
        variant_name = vr.get('variant', 'unknown')
        ats_score = vr.get('ats_score', 0)
        cv_text = vr.get('cv_text_preview', '')[:1500]  # Limit size
        variants_text += f"""
--- {variant_name.upper()} VARIANT (ATS Score: {ats_score:.1f}) ---
{cv_text}
"""
    
    prompt = f"""You are an expert recruiter and CV analyst. Compare these CV variants for the job below and recommend the BEST one.

JOB DESCRIPTION:
{job_text[:2000]}

JOB REQUIREMENTS:
- Title: {job_info.get('job_title', 'Unknown')}
- Company: {job_info.get('company_name', 'Unknown')}
- Required Skills: {', '.join(job_info.get('required_skills', [])[:10])}

CV VARIANTS TO COMPARE:
{variants_text}

Analyze which variant best matches:
1. The job's tone and culture (startup vs corporate)
2. Required skills and keywords
3. Impact/metrics presentation style
4. Overall fit for this specific role

Return ONLY valid JSON:
{{
    "recommended": "professional" or "technical" or "impact",
    "reasoning": "2-3 sentence explanation of why this variant is best for THIS job",
    "confidence": "high" or "medium" or "low"
}}
"""
    
    try:
        response = client.messages.create(
            model=CLAUDE_MODEL_FAST,
            max_tokens=500,
            messages=[{"role": "user", "content": prompt}]
        )
        
        response_text = response.content[0].text
        json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
        
        if json_match:
            result = json.loads(json_match.group())
            print(f"    ✓ LLM recommends: {result.get('recommended')} ({result.get('confidence')} confidence)")
            return result
        
        # Fallback to highest ATS score
        best = max(variant_results, key=lambda x: x.get('ats_score', 0))
        return {
            'recommended': best.get('variant', 'professional'),
            'reasoning': 'Based on highest ATS score.',
            'confidence': 'medium'
        }
        
    except Exception as e:
        print(f"    ⚠ LLM selection failed: {e}, using ATS score")
        best = max(variant_results, key=lambda x: x.get('ats_score', 0))
        return {
            'recommended': best.get('variant', 'professional'),
            'reasoning': f'Fallback to highest ATS score due to error: {str(e)[:50]}',
            'confidence': 'low'
        }


def is_url(text: str) -> bool:
    """Check if text is a valid URL"""
    try:
        result = urlparse(text.strip())
        return all([result.scheme in ['http', 'https'], result.netloc])
    except:
        return False


def parse_urls(input_text: str) -> list:
    """Parse URLs from input text (deduplicates automatically)"""
    urls = []
    seen = set()
    lines = input_text.strip().split('\n')
    
    for line in lines:
        parts = line.split(',')
        for part in parts:
            url = part.strip()
            if url and is_url(url) and url not in seen:
                urls.append(url)
                seen.add(url)
    
    print(f"[DEBUG] parse_urls: Found {len(urls)} unique URLs from {len(lines)} lines")
    return urls


def parse_job_descriptions(input_text: str) -> list:
    """Parse multiple job descriptions from input text (deduplicates automatically)
    
    Only splits on explicit separators (---, ===, etc.) NOT on blank lines.
    Single job postings with formatting blank lines stay as one job.
    """
    # Minimum length for a valid job description (filters out fragments)
    MIN_JOB_LENGTH = 500
    
    # Only use explicit long separators - unlikely to appear in normal text
    separators = ['------------', '============', '************', '############']
    for sep in separators:
        if sep in input_text:
            parts = input_text.split(sep)
            # Filter out short fragments
            jobs = [part.strip() for part in parts if part.strip() and len(part.strip()) >= MIN_JOB_LENGTH]
            if len(jobs) > 1:
                # Deduplicate
                seen = set()
                unique_jobs = []
                for job in jobs:
                    job_hash = hash(job[:200])
                    if job_hash not in seen:
                        unique_jobs.append(job)
                        seen.add(job_hash)
                print(f"[DEBUG] parse_job_descriptions: Found {len(unique_jobs)} jobs using separator '{sep}'")
                return unique_jobs
            elif len(jobs) == 1:
                print(f"[DEBUG] parse_job_descriptions: Single job after filtering (separator '{sep}' present but only 1 substantial section)")
                return jobs
    
    # No explicit separators found - treat entire input as ONE job
    # (Don't split on blank lines - they're just formatting within a single posting)
    print(f"[DEBUG] parse_job_descriptions: Single job description ({len(input_text)} chars, no separators)")
    return [input_text.strip()] if input_text.strip() else []


# Flask Routes

@app.route('/')
def index():
    """Render the main page"""
    return render_template('index.html')


@app.route('/settings')
def settings_page():
    """Render the settings page"""
    settings = get_settings()
    skills = get_all_skills(include_blacklisted=True)
    experiences = get_experiences()
    return render_template('settings.html', settings=settings, skills=skills, experiences=experiences)


@app.route('/api/settings', methods=['GET', 'POST'])
def api_settings():
    """Get or update settings"""
    if request.method == 'GET':
        return jsonify(get_settings())
    
    data = request.json
    for key, value in data.items():
        update_setting(key, value)
    
    return jsonify({'success': True})


@app.route('/api/skills', methods=['GET', 'POST'])
def api_skills():
    """Get all skills or add a new skill"""
    if request.method == 'GET':
        return jsonify(get_all_skills(include_blacklisted=True))
    
    data = request.json
    skill_id = add_skill(
        name=data['name'],
        tier=data.get('tier', 'tier_4_tools'),
        category=data.get('category'),
        is_blacklisted=data.get('is_blacklisted', False)
    )
    
    if skill_id:
        return jsonify({'success': True, 'id': skill_id})
    return jsonify({'success': False, 'error': 'Skill already exists'}), 400


@app.route('/api/skills/<int:skill_id>', methods=['PUT', 'DELETE'])
def api_skill(skill_id):
    """Update or delete a skill"""
    if request.method == 'DELETE':
        delete_skill(skill_id)
        return jsonify({'success': True})
    
    data = request.json
    update_skill(
        skill_id,
        name=data.get('name'),
        tier=data.get('tier'),
        is_blacklisted=data.get('is_blacklisted')
    )
    return jsonify({'success': True})


@app.route('/api/experiences', methods=['GET', 'POST'])
def api_experiences():
    """Get all experiences or add a new one"""
    if request.method == 'GET':
        return jsonify(get_experiences())
    
    data = request.json
    exp_id = add_experience(
        company=data['company'],
        role=data['role'],
        start_date=data.get('start_date'),
        end_date=data.get('end_date'),
        is_current=data.get('is_current', False)
    )
    
    return jsonify({'success': True, 'id': exp_id})


@app.route('/api/experiences/<int:exp_id>', methods=['PUT', 'DELETE'])
def api_experience(exp_id):
    """Update or delete an experience"""
    if request.method == 'DELETE':
        conn = get_db()
        cursor = conn.cursor()
        cursor.execute('DELETE FROM experiences WHERE id = ?', (exp_id,))
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    
    data = request.json
    update_experience(exp_id, **data)
    return jsonify({'success': True})


@app.route('/api/bullet_points', methods=['POST'])
def api_add_bullet():
    """Add a bullet point"""
    data = request.json
    bp_id = add_bullet_point(
        experience_id=data['experience_id'],
        base_description=data['base_description'],
        tech_placeholder=data.get('tech_placeholder')
    )
    return jsonify({'success': True, 'id': bp_id})


@app.route('/api/bullet_points/<int:bp_id>', methods=['PUT', 'DELETE'])
def api_bullet(bp_id):
    """Update or delete a bullet point"""
    if request.method == 'DELETE':
        delete_bullet_point(bp_id)
        return jsonify({'success': True})
    
    data = request.json
    update_bullet_point(
        bp_id,
        base_description=data.get('base_description'),
        tech_placeholder=data.get('tech_placeholder')
    )
    return jsonify({'success': True})


@app.route('/generate', methods=['POST'])
def generate():
    """Generate CVs and cover letters for multiple job descriptions with variants and ATS scoring"""
    if not ANTHROPIC_API_KEY or ANTHROPIC_API_KEY == 'your_anthropic_api_key_here':
        return jsonify({
            'success': False,
            'error': 'Anthropic API key not configured. Please set ANTHROPIC_API_KEY in .env file.'
        }), 400
    
    data = request.json
    input_text = data.get('job_descriptions', '')
    input_mode = data.get('input_mode', 'text')
    generate_variants = data.get('generate_variants', True)  # Generate all 3 variants by default
    
    if not input_text:
        return jsonify({
            'success': False,
            'error': 'No input provided'
        }), 400
    
    # Parse based on input mode
    jobs = []
    scrape_errors = []
    
    print(f"\n[DEBUG] Input mode: {input_mode}")
    print(f"[DEBUG] Input text length: {len(input_text)} chars")
    print(f"[DEBUG] Input preview: {input_text[:200]}...")
    
    if input_mode == 'urls':
        urls = parse_urls(input_text)
        print(f"[DEBUG] Parsed {len(urls)} URLs: {urls}")
        
        if not urls:
            return jsonify({
                'success': False,
                'error': 'No valid URLs found. Please enter URLs starting with http:// or https://'
            }), 400
        
        for url in urls:
            try:
                print(f"[DEBUG] Scraping URL: {url}")
                job_text = scrape_job_url(url)
                print(f"[DEBUG] Scraped {len(job_text)} chars from {url}")
                jobs.append({'text': job_text, 'source_url': url})
            except Exception as e:
                print(f"[DEBUG] Scrape error for {url}: {e}")
                scrape_errors.append({'url': url, 'error': str(e)})
    else:
        parsed = parse_job_descriptions(input_text)
        print(f"[DEBUG] Parsed {len(parsed)} job descriptions from text input")
        for idx, job in enumerate(parsed):
            print(f"[DEBUG] Job {idx+1} length: {len(job)} chars, preview: {job[:80]}...")
        jobs = [{'text': job, 'source_url': None} for job in parsed]
    
    print(f"[DEBUG] Total jobs to process: {len(jobs)}")
    
    if not jobs:
        error_msg = 'Could not parse any job descriptions'
        if scrape_errors:
            error_msg += '. Scraping errors: ' + '; '.join([f"{e['url']}: {e['error']}" for e in scrape_errors])
        return jsonify({
            'success': False,
            'error': error_msg
        }), 400
    
    generator = BatchCVGenerator(ANTHROPIC_API_KEY)
    ats_scorer = ATSScorer()
    results = []
    settings = get_settings()
    
    # Create output directory with readable date/time
    now = datetime.now()
    date_str = now.strftime("%d-%b-%Y")  # e.g., "05-Apr-2026"
    time_str = now.strftime("%H%M")  # e.g., "2058"
    output_dir = Path(f"outputs/{date_str}_{time_str}")
    output_dir.mkdir(parents=True, exist_ok=True)
    
    print(f"\n{'='*60}")
    print(f"STARTING GENERATION: {len(jobs)} job(s) to process")
    print(f"Variants enabled: {generate_variants}")
    print(f"Variants to generate: {list(CV_VARIANTS.keys()) if generate_variants else ['professional']}")
    print(f"{'='*60}\n")
    
    for i, job_data in enumerate(jobs):
        print(f"\n>>> PROCESSING JOB {i+1} of {len(jobs)} <<<")
        job_text = job_data['text']
        source_url = job_data.get('source_url')
        print(f"    Source URL: {source_url}")
        print(f"    Job text length: {len(job_text)} chars")
        print(f"    Job text preview: {job_text[:100]}...")
        
        try:
            # Parse job description
            job_info = generator.parse_job_description(job_text)
            # Store raw text for ATS phrase extraction
            job_info['raw_text'] = job_text
            
            # Create safe filename
            company_safe = re.sub(r'[^\w\s-]', '', job_info.get('company_name', 'Company')).replace(' ', '_')
            title_safe = re.sub(r'[^\w\s-]', '', job_info.get('job_title', 'Position')).replace(' ', '_')
            
            # Create job-specific folder with clear naming
            job_folder = output_dir / f"{company_safe}_{title_safe}"
            job_folder.mkdir(exist_ok=True)
            
            user_name = settings.get('user_name', 'Drew_Gillies').replace(' ', '_')
            job_skills = job_info.get('required_skills', []) + job_info.get('preferred_skills', [])
            
            # Generate variants - filter by feature flags
            if generate_variants:
                variants_to_generate = []
                if ENABLE_PROFESSIONAL:
                    variants_to_generate.append('professional')
                if ENABLE_TECHNICAL:
                    variants_to_generate.append('technical')
                if ENABLE_IMPACT:
                    variants_to_generate.append('impact')
                # Fallback if all disabled
                if not variants_to_generate:
                    variants_to_generate = ['impact']
                    print("    ⚠ All variants disabled in .env, defaulting to impact")
            else:
                variants_to_generate = ['professional']
            variant_results = []
            best_variant = None
            best_ats_score = 0
            
            print(f"\n    --- Generating {len(variants_to_generate)} variants: {variants_to_generate} ---")
            for variant_idx, variant_key in enumerate(variants_to_generate):
                try:
                    print(f"\n  [{variant_idx+1}/{len(variants_to_generate)}] Generating {variant_key} variant...")
                    # Generate CV for this variant
                    generated = generator.generate_cv_and_cover_letter(job_info, variant_key)
                    cv_data = generated.get('cv', {})
                    print(f"    CV data keys: {list(cv_data.keys())}")
                    
                    # Enhance CV content with additional LLM call (if enabled)
                    if ENABLE_ENHANCEMENT:
                        cv_data = llm_enhance_cv_content(cv_data, job_info, variant_key)
                    else:
                        print(f"    ⏭ Skipping enhancement (disabled in .env)")
                    
                    # LLM-as-judge: Opus reviews and fixes the final CV
                    if ENABLE_JUDGE:
                        cv_data = llm_judge_cv_feedback(cv_data, job_info)
                    else:
                        print(f"    ⏭ Skipping judge (disabled in .env)")
                    
                    # Build CV text for ATS scoring
                    cv_text = f"""
                    {cv_data.get('bio', '')}
                    Skills: {', '.join(cv_data.get('expertise', []))}
                    Experience:
                    {cv_data.get('c', {}).get('skills', '')}
                    {cv_data.get('c', {}).get('bp1', '')}
                    {cv_data.get('c', {}).get('bp2', '')}
                    {cv_data.get('c', {}).get('bp3', '')}
                    {cv_data.get('c', {}).get('bp4', '')}
                    {cv_data.get('t', {}).get('skills', '')}
                    {cv_data.get('t', {}).get('bp1', '')}
                    {cv_data.get('t', {}).get('bp2', '')}
                    {cv_data.get('t', {}).get('bp3', '')}
                    {cv_data.get('t', {}).get('bp4', '')}
                    {cv_data.get('a', {}).get('skills', '')}
                    {cv_data.get('a', {}).get('bp1', '')}
                    {cv_data.get('a', {}).get('bp2', '')}
                    {cv_data.get('a', {}).get('bp3', '')}
                    """
                    
                    # Score with ATS
                    ats_result = ats_scorer.score_cv(cv_text, job_text, job_skills)
                    
                    # Track best variant
                    if ats_result['total_score'] > best_ats_score:
                        best_ats_score = ats_result['total_score']
                        best_variant = variant_key
                    
                    # Save CV (Word + PDF)
                    variant_name = CV_VARIANTS[variant_key]['name']
                    cv_docx_path = job_folder / f"CV_{variant_name}_{company_safe}.docx"
                    cv_pdf_path = generator.create_cv_docx(cv_data, job_info, str(cv_docx_path), create_pdf=True)
                    
                    # Extract ATS feedback for UI
                    kw_details = ats_result.get('breakdown', {}).get('keywords', {}).get('details', {})
                    missing_phrases = kw_details.get('missing_phrases', [])[:5]
                    missing_keywords = kw_details.get('missing_keywords', [])[:3]
                    
                    print(f"    ✓ {variant_key} complete - ATS score: {ats_result['total_score']:.1f}")
                    variant_results.append({
                        'variant': variant_key,
                        'variant_name': variant_name,
                        'cv_docx_path': str(cv_docx_path),
                        'cv_pdf_path': str(cv_pdf_path) if cv_pdf_path else None,
                        'ats_score': ats_result['total_score'],
                        'ats_grade': ats_result['grade'],
                        'ats_pass_likelihood': ats_result['ats_pass_likelihood'],
                        'ats_feedback': ats_result.get('feedback', []),
                        'missing_phrases': missing_phrases,
                        'missing_keywords': missing_keywords,
                        'page_valid': generated.get('page_validation', {}).get('is_valid', True),
                        'page_estimate': generated.get('page_validation', {}).get('estimated_pages', 1.0),
                        'cv_text_preview': cv_text,  # For LLM comparison
                        'cover_letter': generated.get('cover_letter', ''),  # Reuse later, no extra LLM call
                    })
                    
                except Exception as e:
                    import traceback
                    print(f"    ✗ Error generating {variant_key}: {e}")
                    traceback.print_exc()
                    variant_results.append({
                        'variant': variant_key,
                        'error': str(e)
                    })
            
            print(f"\n    --- All {len(variant_results)} variants complete for job {i+1} ---")
            print(f"    Best variant (ATS): {best_variant} with score: {best_ats_score:.1f}")
            
            # LLM selection of best variant (if multiple variants generated)
            llm_recommendation = None
            successful_variants = [v for v in variant_results if not v.get('error')]
            if len(successful_variants) >= 2:
                print(f"    Asking LLM to select best variant...")
                llm_recommendation = llm_select_best_variant(successful_variants, job_text, job_info)
                # Mark the recommended variant
                for vr in variant_results:
                    vr['llm_recommended'] = (vr.get('variant') == llm_recommendation.get('recommended'))
                    if vr['llm_recommended']:
                        vr['llm_reasoning'] = llm_recommendation.get('reasoning', '')
                        vr['llm_confidence'] = llm_recommendation.get('confidence', 'medium')
            else:
                # Only one variant, mark it as recommended
                for vr in variant_results:
                    vr['llm_recommended'] = True
                    vr['llm_reasoning'] = 'Only variant generated.'
                    vr['llm_confidence'] = 'high'
            
            # Generate cover letter (if enabled)
            cover_letter_path = None
            if ENABLE_COVER_LETTER:
                # Reuse the cover letter already produced during variant generation
                # (generate_cv_and_cover_letter returns both) - avoids a second full LLM call.
                best_vr = next(
                    (vr for vr in variant_results
                     if vr.get('variant') == best_variant and not vr.get('error')),
                    None
                )
                cover_letter_text = best_vr.get('cover_letter', '') if best_vr else ''
                # Fall back to the first variant that produced a cover letter
                if not cover_letter_text:
                    cover_letter_text = next(
                        (vr.get('cover_letter', '') for vr in variant_results
                         if vr.get('cover_letter')),
                        ''
                    )

                if cover_letter_text:
                    print(f"    Writing cover letter (reused from {best_variant or 'first'} variant)...")
                    cover_letter_path = job_folder / f"Cover_Letter_{company_safe}.pdf"
                    generator.create_cover_letter_pdf(
                        cover_letter_text,
                        job_info,
                        str(cover_letter_path)
                    )
                    print(f"    ✓ Cover letter saved to {cover_letter_path}")
                else:
                    print(f"    ⚠ No cover letter available from variants, skipping")
            else:
                print(f"    ⏭ Skipping cover letter (disabled in .env)")
            print(f"\n>>> JOB {i+1} COMPLETE <<<")
            
            # Save original job description
            job_desc_path = job_folder / "Original_Job_Description.txt"
            with open(job_desc_path, 'w', encoding='utf-8') as f:
                f.write(f"Job Title: {job_info.get('job_title', 'Unknown')}\n")
                f.write(f"Company: {job_info.get('company_name', 'Unknown')}\n")
                f.write(f"Date Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                if source_url:
                    f.write(f"Source URL: {source_url}\n")
                f.write(f"\nBest Variant: {best_variant} (ATS Score: {best_ats_score:.1f})\n")
                f.write(f"\n{'='*50}\n")
                f.write("ORIGINAL JOB DESCRIPTION:\n")
                f.write(f"{'='*50}\n\n")
                f.write(job_text)
            
            # Save ATS report
            ats_report_path = job_folder / "ATS_Score_Report.txt"
            with open(ats_report_path, 'w', encoding='utf-8') as f:
                f.write("ATS COMPATIBILITY REPORT\n")
                f.write("=" * 50 + "\n\n")
                for vr in variant_results:
                    if 'error' not in vr:
                        f.write(f"{vr['variant_name']} Variant:\n")
                        f.write(f"  Score: {vr['ats_score']:.1f}/100 ({vr['ats_grade']})\n")
                        f.write(f"  Pass Likelihood: {vr['ats_pass_likelihood']}\n")
                        f.write(f"  Page Valid: {'Yes' if vr['page_valid'] else 'No'} ({vr['page_estimate']:.2f} pages)\n\n")
                f.write(f"\nRECOMMENDED: {CV_VARIANTS.get(best_variant, {}).get('name', 'Professional')} variant\n")
            
            results.append({
                'success': True,
                'job_title': job_info.get('job_title', 'Unknown'),
                'company': job_info.get('company_name', 'Unknown'),
                'source_url': source_url,
                'variants': variant_results,
                'best_variant': best_variant,
                'best_ats_score': best_ats_score,
                'cover_letter_path': str(cover_letter_path) if cover_letter_path else None,
                'folder': str(job_folder)
            })
            
        except Exception as e:
            import traceback
            print(f"\n!!! EXCEPTION processing job {i+1} !!!")
            print(f"Error: {e}")
            traceback.print_exc()
            results.append({
                'success': False,
                'error': str(e),
                'source_url': source_url,
                'job_text_preview': job_text[:100] + '...' if len(job_text) > 100 else job_text
            })
    
    print(f"\n{'='*60}")
    print(f"ALL JOBS COMPLETE!")
    print(f"  Processed: {len(jobs)} job(s)")
    print(f"  Successful: {sum(1 for r in results if r.get('success'))}")
    print(f"  Failed: {sum(1 for r in results if not r.get('success'))}")
    print(f"{'='*60}\n")
    
    # Create zip file of all outputs
    print("Creating zip file...")
    zip_path = output_dir / "all_applications.zip"
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for result in results:
            if result.get('success'):
                folder = Path(result['folder'])
                for file in folder.iterdir():
                    zipf.write(file, f"{folder.name}/{file.name}")
    
    print(f"✓ Zip file created: {zip_path}")
    print("Returning JSON response to client...\n")
    
    return jsonify({
        'success': True,
        'results': results,
        'output_directory': str(output_dir),
        'zip_file': str(zip_path),
        'total_jobs': len(jobs),
        'successful': sum(1 for r in results if r.get('success')),
        'failed': sum(1 for r in results if not r.get('success')),
        'scrape_errors': scrape_errors,
        'variants_generated': list(CV_VARIANTS.keys()) if generate_variants else ['professional']
    })


@app.route('/download/<path:filepath>')
def download(filepath):
    """Download a generated file"""
    try:
        return send_file(filepath, as_attachment=True)
    except Exception as e:
        return jsonify({'error': str(e)}), 404


if __name__ == '__main__':
    Path('templates').mkdir(exist_ok=True)
    
    print("🚀 Starting CV Generator Web App")
    print("=" * 40)
    
    # Show testing flags
    print("\n� Testing Flags (.env):")
    print(f"   ENABLE_PROFESSIONAL: {ENABLE_PROFESSIONAL}")
    print(f"   ENABLE_TECHNICAL:    {ENABLE_TECHNICAL}")
    print(f"   ENABLE_IMPACT:       {ENABLE_IMPACT}")
    print(f"   ENABLE_COVER_LETTER: {ENABLE_COVER_LETTER}")
    print(f"   ENABLE_ENHANCEMENT:  {ENABLE_ENHANCEMENT}")
    
    enabled_variants = [v for v, e in [('professional', ENABLE_PROFESSIONAL), ('technical', ENABLE_TECHNICAL), ('impact', ENABLE_IMPACT)] if e]
    print(f"\n   → Will generate: {enabled_variants or ['impact (fallback)']}")
    
    if not ANTHROPIC_API_KEY or ANTHROPIC_API_KEY == 'your_anthropic_api_key_here':
        print("\n⚠️  Warning: ANTHROPIC_API_KEY not set in .env file")
        print("   Create a .env file with: ANTHROPIC_API_KEY=your_key_here")
    else:
        print("\n✓ Anthropic API key loaded")
    
    print("\n📝 Open http://localhost:5001 in your browser")
    print("⚙️  Settings page: http://localhost:5001/settings")
    app.run(debug=True, port=5001)
