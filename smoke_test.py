#!/usr/bin/env python3
"""
Smoke tests for CV Generator - tests end-to-end functionality
"""

import os
import sys
from pathlib import Path

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent))

from dotenv import load_dotenv
load_dotenv()

def test_template_exists():
    """Test that the CV template file exists"""
    template_path = 'resources/template.docx'
    exists = Path(template_path).exists()
    print(f"✓ Template exists: {template_path}" if exists else f"✗ Template MISSING: {template_path}")
    return exists

def test_database_connection():
    """Test database connection and data retrieval"""
    try:
        from database import get_settings, get_experiences, get_all_skills
        settings = get_settings()
        experiences = get_experiences()
        skills = get_all_skills()
        print(f"✓ Database connected - {len(settings)} settings, {len(experiences)} experiences, {len(skills)} skills")
        return True
    except Exception as e:
        print(f"✗ Database error: {e}")
        return False

def test_api_key():
    """Test that Anthropic API key is configured"""
    api_key = os.getenv('ANTHROPIC_API_KEY')
    if api_key and api_key != 'your_anthropic_api_key_here':
        print(f"✓ API key configured (starts with: {api_key[:10]}...)")
        return True
    else:
        print("✗ API key NOT configured")
        return False

def test_cv_generation():
    """Test CV document generation from template"""
    from docx import Document
    from pathlib import Path
    
    template_path = 'resources/template.docx'
    if not Path(template_path).exists():
        print("✗ Cannot test CV generation - template missing")
        return False
    
    try:
        doc = Document(template_path)
        
        # Check for placeholders in template
        placeholders_found = []
        for para in doc.paragraphs:
            if '{{' in para.text:
                placeholders_found.append(para.text[:50])
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if '{{' in para.text:
                            placeholders_found.append(para.text[:50])
        
        if placeholders_found:
            print(f"✓ Template has {len(placeholders_found)} placeholders")
            for p in placeholders_found[:5]:
                print(f"    - {p}...")
        else:
            print("⚠ No placeholders found in template (may use different format)")
        
        # Test saving
        test_output = Path('outputs/test_cv.docx')
        test_output.parent.mkdir(exist_ok=True)
        doc.save(str(test_output))
        
        if test_output.exists():
            print(f"✓ Can save DOCX to: {test_output}")
            test_output.unlink()  # Clean up
            return True
        else:
            print("✗ Failed to save test DOCX")
            return False
            
    except Exception as e:
        print(f"✗ CV generation error: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_ats_scorer():
    """Test ATS scoring functionality"""
    try:
        from ats_scorer import ATSScorer, validate_single_page
        
        scorer = ATSScorer()
        
        # Test with sample data
        cv_text = """
        Software Engineer with 3 years experience in Python, AWS, and cloud infrastructure.
        Skills: Python, Java, AWS, Docker, Kubernetes, PostgreSQL
        Experience:
        - Built scalable microservices handling 10,000 requests per second
        - Reduced deployment time by 50% through CI/CD automation
        - Led team of 5 engineers on cloud migration project
        """
        
        job_text = """
        We are looking for a Software Engineer with Python and AWS experience.
        Requirements: Python, AWS, Docker, CI/CD
        """
        
        job_skills = ['Python', 'AWS', 'Docker', 'CI/CD']
        
        result = scorer.score_cv(cv_text, job_text, job_skills)
        
        print(f"✓ ATS Scorer works - Score: {result['total_score']:.1f}/100 ({result['grade']})")
        return True
        
    except Exception as e:
        print(f"✗ ATS Scorer error: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_full_generation():
    """Test full CV generation with Claude API"""
    api_key = os.getenv('ANTHROPIC_API_KEY')
    if not api_key or api_key == 'your_anthropic_api_key_here':
        print("⚠ Skipping full generation test - no API key")
        return None
    
    try:
        from web_app import BatchCVGenerator, CV_VARIANTS
        from pathlib import Path
        
        generator = BatchCVGenerator(api_key)
        
        # Test job parsing
        test_job = """
        Software Engineer at TechCorp
        Location: London, UK
        
        We're looking for a Python developer with AWS experience.
        Requirements:
        - 2+ years Python
        - AWS (Lambda, S3, EC2)
        - Docker
        - PostgreSQL
        
        Nice to have:
        - Kubernetes
        - CI/CD experience
        """
        
        print("  Testing job parsing...")
        job_info = generator.parse_job_description(test_job)
        print(f"    Parsed: {job_info.get('job_title')} at {job_info.get('company_name')}")
        
        print("  Testing CV generation (professional variant)...")
        generated = generator.generate_cv_and_cover_letter(job_info, 'professional')
        
        cv_data = generated.get('cv', {})
        if cv_data.get('bio'):
            print(f"    Bio: {cv_data['bio'][:80]}...")
        if cv_data.get('expertise'):
            print(f"    Expertise: {len(cv_data['expertise'])} skills")
        
        # Test document creation
        print("  Testing DOCX creation...")
        test_output = Path('outputs/smoke_test')
        test_output.mkdir(parents=True, exist_ok=True)
        
        cv_path = test_output / "test_cv.docx"
        pdf_path = generator.create_cv_docx(cv_data, job_info, str(cv_path), create_pdf=True)
        
        if cv_path.exists():
            print(f"✓ DOCX created: {cv_path} ({cv_path.stat().st_size} bytes)")
        else:
            print(f"✗ DOCX NOT created: {cv_path}")
            return False
        
        if pdf_path and Path(pdf_path).exists():
            print(f"✓ PDF created: {pdf_path}")
        else:
            print(f"⚠ PDF not created (converter may not be available)")
        
        # Clean up
        if cv_path.exists():
            cv_path.unlink()
        if pdf_path and Path(pdf_path).exists():
            Path(pdf_path).unlink()
        
        print("✓ Full generation test PASSED")
        return True
        
    except Exception as e:
        print(f"✗ Full generation error: {e}")
        import traceback
        traceback.print_exc()
        return False

def run_all_tests():
    """Run all smoke tests"""
    print("=" * 60)
    print("CV GENERATOR SMOKE TESTS")
    print("=" * 60)
    
    results = {}
    
    print("\n1. Template File")
    results['template'] = test_template_exists()
    
    print("\n2. Database Connection")
    results['database'] = test_database_connection()
    
    print("\n3. API Key")
    results['api_key'] = test_api_key()
    
    print("\n4. CV Document Generation")
    results['cv_gen'] = test_cv_generation()
    
    print("\n5. ATS Scorer")
    results['ats'] = test_ats_scorer()
    
    print("\n6. Full End-to-End Generation")
    results['full'] = test_full_generation()
    
    print("\n" + "=" * 60)
    print("SUMMARY")
    print("=" * 60)
    
    passed = sum(1 for v in results.values() if v is True)
    failed = sum(1 for v in results.values() if v is False)
    skipped = sum(1 for v in results.values() if v is None)
    
    print(f"Passed: {passed}")
    print(f"Failed: {failed}")
    print(f"Skipped: {skipped}")
    
    if failed > 0:
        print("\n⚠ Some tests FAILED - check output above")
        return False
    else:
        print("\n✓ All tests PASSED")
        return True

if __name__ == '__main__':
    success = run_all_tests()
    sys.exit(0 if success else 1)
