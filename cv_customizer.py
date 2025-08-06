import json
import re
from pathlib import Path
from typing import Dict, List, Union
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import subprocess
import sys
from datetime import datetime
import os

class CVCustomizer:
    def __init__(self, template_path: str):
        """initialize with a word template containing placeholders"""
        self.template_path = template_path
        self.document = Document(template_path)
    
    def replace_placeholders(self, replacements: Dict[str, Union[str, List[str]]]):
        """replace placeholders throughout the document while preserving formatting"""
        # handle paragraphs
        for paragraph in self.document.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)
        
        # handle tables
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)
        
        # handle headers and footers
        for section in self.document.sections:
            # header
            for paragraph in section.header.paragraphs:
                self._replace_in_paragraph(paragraph, replacements)
            # footer
            for paragraph in section.footer.paragraphs:
                self._replace_in_paragraph(paragraph, replacements)
    
# In cv_customizer.py

    def _replace_in_paragraph(self, paragraph, replacements):
        """replace placeholders in a paragraph while preserving run formatting"""
        full_text = paragraph.text
        
        flat_replacements = {}
        for key, value in replacements.items():
            if isinstance(value, dict):
                for subkey, subvalue in value.items():
                    flat_replacements[f"{key}.{subkey}"] = subvalue
            else:
                flat_replacements[key] = value
        
        has_placeholder = any(f"{{{{{key}}}}}" in full_text for key in flat_replacements)
        if not has_placeholder:
            return
        
        for key, value in flat_replacements.items():
            placeholder = f"{{{{{key}}}}}"
            
            if placeholder in full_text:
                value_text = ""
                # --- START OF NEW LOGIC ---
                # Special handling for tech stack keys to ensure they are a single, comma-separated string.
                if key in ['t.skills', 'a.skills']:
                    if isinstance(value, list):
                        value_text = ', '.join(value)
                    else:
                        value_text = str(value) # It's already a string, use as is.
                # For all other lists (like 'expertise'), create bullet points.
                elif isinstance(value, list):
                    value_text = '\n• '.join(value)
                    value_text = '• ' + value_text if value else ''
                # Handle all other non-list values.
                else:
                    value_text = str(value).replace('\n', ' ').strip()
                # --- END OF NEW LOGIC ---

                # If it's a simple replacement (placeholder is complete in one run)
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value_text)
                        return
                
                # Complex case: placeholder spans multiple runs
                self._complex_replace(paragraph, placeholder, value_text)
    
    def _complex_replace(self, paragraph, placeholder, replacement):
        """handle placeholders that span multiple runs"""
        # store run properties
        run_props = []
        combined_text = ""
        
        for run in paragraph.runs:
            run_props.append({
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'color': run.font.color.rgb if run.font.color.rgb else None
            })
            combined_text += run.text
        
        # replace in combined text
        new_text = combined_text.replace(placeholder, replacement)
        
        # clear runs and recreate with original formatting
        for run in paragraph.runs:
            run.text = ""
        
        # add the new text with the first run's formatting
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            # apply the original formatting from the first run
            if run_props:
                props = run_props[0]
                paragraph.runs[0].bold = props['bold']
                paragraph.runs[0].italic = props['italic']
                paragraph.runs[0].underline = props['underline']
                if props['font_name']:
                    paragraph.runs[0].font.name = props['font_name']
                if props['font_size']:
                    paragraph.runs[0].font.size = props['font_size']
    
    def save_docx(self, output_path: str):
        """save the modified document as a word file"""
        self.document.save(output_path)
    
    def convert_to_pdf(self, docx_path: str, pdf_path: str):
        """convert word document to pdf using libreoffice (cross-platform)"""
        try:
            # using libreoffice in headless mode
            subprocess.run([
                'soffice',
                '--headless',
                '--convert-to',
                'pdf',
                '--outdir',
                str(Path(pdf_path).parent),
                docx_path
            ], check=True, capture_output=True)
            
            # rename to desired output name if needed
            generated_pdf = Path(docx_path).with_suffix('.pdf')
            if generated_pdf.name != Path(pdf_path).name:
                generated_pdf.rename(pdf_path)
                
        except FileNotFoundError:
            raise Exception("LibreOffice not found. Install with: brew install --cask libreoffice")
        except subprocess.CalledProcessError as e:
            # fallback to python-docx2pdf (windows/mac only)
            try:
                from docx2pdf import convert
                convert(docx_path, pdf_path)
                print("✓ Used docx2pdf as fallback")
            except ImportError:
                raise Exception("PDF conversion failed. Install LibreOffice: brew install --cask libreoffice")
            except Exception as docx2pdf_error:
                raise Exception(f"PDF conversion failed: {docx2pdf_error}")
        except Exception as e:
            raise Exception(f"PDF conversion failed: {e}")
    
    def customize_cv(self, job_data: Dict, output_name: str, job_info=None, original_job_text: str = ""):
        """main method to customize cv with job data"""
        # Validate content length to ensure 1-page format
        self._validate_content_length(job_data)
        
        # Debug: Print the data structure being passed
        print("🔧 CV data structure:")
        for key, value in job_data.items():
            if isinstance(value, dict):
                print(f"  {key}:")
                for subkey, subvalue in value.items():
                    print(f"    {subkey}: {str(subvalue)[:60]}...")
            else:
                print(f"  {key}: {str(value)[:60]}...")
        
        # Create improved execution folder structure
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Clean company and role names for folder
        if job_info:
            company_clean = re.sub(r'[^\w\s-]', '', job_info.company_name or 'Unknown_Company').replace(' ', '_')
            role_clean = re.sub(r'[^\w\s-]', '', job_info.job_title or 'Software_Role').replace(' ', '_')
            execution_folder = f"outputs/{timestamp}-{company_clean}-{role_clean}"
        else:
            execution_folder = f"outputs/{timestamp}-{output_name}"
        
        # create the execution folder
        os.makedirs(execution_folder, exist_ok=True)
        print(f"📁 Created execution folder: {execution_folder}")
        
        # Save original job description
        if original_job_text:
            job_desc_path = f"{execution_folder}/Original_Job_Description.txt"
            with open(job_desc_path, 'w', encoding='utf-8') as f:
                f.write(f"Job Title: {job_info.job_title if job_info else 'Unknown'}\n")
                f.write(f"Company: {job_info.company_name if job_info else 'Unknown'}\n")
                f.write(f"Date Applied: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Location: {job_info.location if job_info else 'Unknown'}\n")
                f.write("\n" + "="*50 + "\n")
                f.write("ORIGINAL JOB DESCRIPTION:\n")
                f.write("="*50 + "\n\n")
                f.write(original_job_text)
            print(f"💾 Saved job description: {job_desc_path}")
        
        # replace all placeholders
        self.replace_placeholders(job_data)
        
        # save as word document with standard name
        docx_output = f"{execution_folder}/Drew_Gillies_Software_Resume.docx"
        self.save_docx(docx_output)
        print(f"✓ Word document saved: {docx_output}")
        
        # convert to pdf
        pdf_output = f"{execution_folder}/Drew_Gillies_Software_Resume.pdf"
        try:
            self.convert_to_pdf(docx_output, pdf_output)
            print(f"✓ PDF generated: {pdf_output}")
        except Exception as e:
            print(f"⚠️  PDF conversion failed: {e}")
            print("   Word document is still available")
            pdf_output = None
        
        return docx_output, pdf_output, execution_folder
    
# In cv_customizer.py

    def estimate_page_usage(self, job_data: Dict) -> float:
        """Estimate how much of the page the CV will use (0.0 to 1.0) - FINAL CALIBRATION"""

        # Recalibrated weights for more accurate estimation
        char_weights = {
            'bio': 1.8,
            'expertise_skill': 10,
            'tech_stack': 1.5,
            'bullet_point': 2.2,
            'section_headers': 40
        }

        total_weighted_chars = 0

        # Count bio characters
        bio = job_data.get('bio', '')
        total_weighted_chars += len(bio) * char_weights['bio']

        # Count expertise skills
        expertise = job_data.get('expertise', [])
        expertise_count = len(expertise) if isinstance(expertise, list) else 0
        total_weighted_chars += expertise_count * char_weights['expertise_skill']

        # Add section headers and other fixed spacing
        total_weighted_chars += char_weights['section_headers'] * 3

        # Count bullet points and tech stacks
        for section_key in ['t', 'a']:
            if section_key in job_data and isinstance(job_data[section_key], dict):
                section_data = job_data[section_key]

                # Tech stack
                skills = section_data.get('skills', '')
                total_weighted_chars += len(skills) * char_weights['tech_stack']

                # Bullet points
                bullet_count = 0
                for i in range(1, 5):
                    bp_key = f'bp{i}'
                    if bp_key in section_data:
                        bp_text = section_data.get(bp_key, '')
                        total_weighted_chars += len(bp_text) * char_weights['bullet_point']
                        bullet_count += 1
                
                # Add overhead for each bullet
                total_weighted_chars += bullet_count * 25

        # A fully packed page on your template holds around 7500 weighted characters.
        # This is the new "magic number" based on the recalibrated weights.
        PAGE_CAPACITY = 7500.0 
        page_usage = total_weighted_chars / PAGE_CAPACITY

        print("📊 Final Page Estimation:")
        print(f"   Total weighted chars: {total_weighted_chars:.0f}")
        print(f"   Final Page Capacity: {PAGE_CAPACITY:.0f}")
        print(f"   Estimated Page Usage: {page_usage:.1%}")

        return min(page_usage, 1.5)
    
    def _validate_content_length(self, job_data: Dict):
        """Validate that content will fit on one page"""
        warnings = []
        
        # Check bio length
        bio = job_data.get('bio', '')
        if len(bio) > 600:  # Increased from 400 for longer bios
            warnings.append(f"Bio is too long ({len(bio)} chars, recommend <600)")
        
        # Skip bullet point validation - now handled by length enforcement system
        # The new system requires 65-75 word bullet points, so character limits are obsolete
        
        if warnings:
            print("⚠️  Content length warnings:")
            for warning in warnings:
                print(f"   - {warning}")
        else:
            print("✓ Content length validation passed")